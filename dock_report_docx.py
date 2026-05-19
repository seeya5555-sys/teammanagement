"""
Dry Dock Report — Word(.docx) 생성

첨부 양식 (PACIFIC BUSAN호) 구조를 그대로 재현:
  · 표지: 보고서 제목 + 선박 정보 표 + 결재란
  · 본문: 1단계 큰 제목 → 1) 2) 3) 하위 → 본문 (불릿/표/사진)
  · 1단계와 1단계 사이는 빈 줄로만 구분 (페이지 분리 X)
  · 사용된 마커: 1단계="1. 2.", 2단계="1) 2)", 불릿마커는 원본 그대로(•, –, 1), a))
"""
import io
import json
import os
from typing import Dict, List, Tuple
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────────────────
#  XML 유틸 (셀 음영 / 페이지 설정 등)
# ─────────────────────────────────────────────────────────────
def _set_cell_shading(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)


def _set_cell_borders(cell, color='808080', sz=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for side in ['top', 'left', 'bottom', 'right']:
        b = tcBorders.find(qn(f'w:{side}'))
        if b is None:
            b = OxmlElement(f'w:{side}')
            tcBorders.append(b)
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(sz))
        b.set(qn('w:color'), color)


def _set_font(run, *, name='Malgun Gothic', size=10, bold=False, color=None):
    run.font.name = name
    # 한글 폰트도 함께 지정
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color.lstrip('#'))


def _add_paragraph(doc_or_cell, text='', *, font='Malgun Gothic', size=10, bold=False,
                   color=None, align=None, before=0, after=0, indent_left=0):
    p = doc_or_cell.add_paragraph()
    pf = p.paragraph_format
    if align is not None:
        p.alignment = align
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if indent_left:
        pf.left_indent = Cm(indent_left)
    if text:
        r = p.add_run(text)
        _set_font(r, name=font, size=size, bold=bold, color=color)
    return p


# ─────────────────────────────────────────────────────────────
#  표지 — 결재란 + 보고서 제목 + 선박 정보 표
# ─────────────────────────────────────────────────────────────
def _build_cover(doc, report):
    """첨부 양식의 표지 페이지"""

    # ① 결재란 (우측 정렬) — 단일 표
    approvals = [
        ('기 안', report.get('approval_drafter') or ''),
        ('팀 장', report.get('approval_team_lead') or ''),
        ('중 역', report.get('approval_director') or ''),
        ('대표이사', report.get('approval_ceo') or ''),
    ]

    # 결재 라벨 1열 + 결재자 4열 = 총 5열, 2행
    approval_tbl = doc.add_table(rows=2, cols=len(approvals) + 1)
    approval_tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
    approval_tbl.autofit = False

    # 표 자체에 left indent를 줘서 우측으로 밀어붙임
    # 본문폭 17cm, 표 총 9.4cm → 좌측 들여쓰기 ~7.5cm
    tblPr = approval_tbl._element.find(qn('w:tblPr'))
    if tblPr is not None:
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), '4250')  # twips (약 7.5cm)
        tblInd.set(qn('w:type'), 'dxa')
        # tblPr 안에서 적절한 위치에 삽입
        existing = tblPr.find(qn('w:tblInd'))
        if existing is not None:
            tblPr.remove(existing)
        tblPr.append(tblInd)

    col_widths_cm = [1.4] + [2.0] * len(approvals)
    # 첫 열: "결재" 라벨 (세로 병합)
    title_cell = approval_tbl.rows[0].cells[0]
    title_cell.merge(approval_tbl.rows[1].cells[0])
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    title_cell.width = Cm(col_widths_cm[0])
    _set_cell_shading(title_cell, 'F2F2F2')
    # "결재" 한 줄로 (Word에서 잘 보이도록)
    p = title_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('결재')
    _set_font(r, size=11, bold=True)
    _set_cell_borders(title_cell)

    # 결재자 헤더 + 빈 서명란
    for col, (label, name) in enumerate(approvals, start=1):
        h = approval_tbl.rows[0].cells[col]
        h.width = Cm(col_widths_cm[col])
        _set_cell_shading(h, 'F2F2F2')
        h.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        ph = h.paragraphs[0]
        ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rh = ph.add_run(label)
        _set_font(rh, size=9, bold=True)
        _set_cell_borders(h)

        sig = approval_tbl.rows[1].cells[col]
        sig.width = Cm(col_widths_cm[col])
        sig.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        ps = sig.paragraphs[0]
        ps.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if name:
            rs = ps.add_run(name)
            _set_font(rs, size=10)
        _set_cell_borders(sig)

    # 서명란 행 높이 고정
    _set_row_height(approval_tbl.rows[1], 1.3)

    _add_paragraph(doc, '', before=12)

    # ② 보고서 제목
    title = report.get('title') or 'Dry Dock Report'
    _add_paragraph(doc, title,
                   size=18, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   before=24, after=8)

    # 제목 밑줄선
    line_p = doc.add_paragraph()
    pPr = line_p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '18')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F4E79')
    pBdr.append(bottom)
    pPr.append(pBdr)

    _add_paragraph(doc, '', before=8)

    # ③ 선박 정보 표 (라벨 | 값 | 라벨 | 값) - 2행 4열
    info_rows = [
        ('Vessel Name', report.get('vessel_name') or '',
         'Type', report.get('vessel_type') or ''),
        ('IMO No.',    report.get('imo_no') or '',
         'Built',      ''),  # built 정보는 양식에 없으나 자리 유지
        ('Gross Tonnage', report.get('gross_tonnage') or '',
         'Dead Weight',   report.get('dead_weight') or ''),
        ('Shipyard',  report.get('shipyard') or '',
         'Dock No.',  report.get('dock_no') or ''),
        ('Dry Dock Period',
         _fmt_period(report.get('period_start'), report.get('period_end')),
         'Reported on', datetime.now().strftime('%Y-%m-%d')),
    ]

    info_tbl = doc.add_table(rows=len(info_rows), cols=4)
    info_tbl.autofit = False
    widths = [Cm(3.5), Cm(5.5), Cm(3.5), Cm(5.5)]
    for ri, row in enumerate(info_rows):
        for ci, val in enumerate(row):
            cell = info_tbl.rows[ri].cells[ci]
            cell.width = widths[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_borders(cell)
            p = cell.paragraphs[0]
            # 라벨 셀(짝수 인덱스): 음영 + 굵게
            if ci % 2 == 0:
                _set_cell_shading(cell, 'E7EBF0')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(val)
                _set_font(r, size=10, bold=True)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(val)
                _set_font(r, size=10)

    # 페이지 나누기 (본문은 새 페이지부터)
    doc.add_page_break()


def _set_no_border(cell, side):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    b = tcBorders.find(qn(f'w:{side}'))
    if b is None:
        b = OxmlElement(f'w:{side}')
        tcBorders.append(b)
    b.set(qn('w:val'), 'nil')


def _set_row_height(row, cm):
    """행 높이 강제 설정 (cm)"""
    tr = row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(cm * 567)))  # 1cm = 567 twips
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


def _fmt_period(start, end):
    if not start and not end:
        return ''
    s = (start or '').replace('-', '.')
    e = (end or '').replace('-', '.')
    if s and e:
        try:
            d1 = datetime.strptime(start, '%Y-%m-%d')
            d2 = datetime.strptime(end, '%Y-%m-%d')
            days = (d2 - d1).days + 1
            return f'{s} ~ {e}  ({days}일)'
        except Exception:
            return f'{s} ~ {e}'
    return s or e


# ─────────────────────────────────────────────────────────────
#  목차 (Table of Contents) — 자동 생성, 페이지 번호 없이 트리 형태
# ─────────────────────────────────────────────────────────────
def _build_toc(doc, sections_tree):
    _add_paragraph(doc, '목 차',
                   size=16, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER,
                   before=12, after=18)

    # tree 평면화 + 번호링
    def walk(nodes, prefix='', depth=0):
        for i, n in enumerate(nodes):
            num = f'{prefix}-{i + 1}' if prefix else f'{i + 1}'
            yield depth, num, n
            yield from walk(n.get('children', []), num, depth + 1)

    for depth, num, n in walk(sections_tree):
        if depth == 0:
            text = f'{num}.  {n["title"]}'
            _add_paragraph(doc, text, size=11.5, bold=True,
                           before=4, after=2)
        elif depth == 1:
            text = f'{num.split("-", 1)[1]})  {n["title"]}'  # "1-2" → "2)"
            _add_paragraph(doc, text, size=10.5, indent_left=0.8,
                           before=2, after=1)
        else:
            text = f'- {n["title"]}'
            _add_paragraph(doc, text, size=10, indent_left=1.6,
                           before=1, after=1)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────
#  본문 — 섹션 트리 + 블록 렌더링
# ─────────────────────────────────────────────────────────────
def _render_sections(doc, sections_tree, depth=0, prefix=''):
    """재귀적으로 섹션과 그 아래 블록 + 자식 섹션 렌더링"""
    for i, sec in enumerate(sections_tree):
        num = f'{prefix}-{i + 1}' if prefix else f'{i + 1}'

        # 1단계가 두 번째 이상이면 새 페이지로 시작
        if depth == 0 and i > 0:
            doc.add_page_break()

        if depth == 0:
            # 1단계 — 큰 제목 ("1.")
            _add_paragraph(doc, f'{num}. {sec["title"]}',
                           size=14, bold=True, color='1F4E79',
                           before=18, after=8)
            # 제목 밑 선
            _add_horizontal_line(doc, color='1F4E79', sz=8)
        elif depth == 1:
            # 2단계 — 하위 제목 ("1) ...")
            sub_num = num.split('-', 1)[1]
            _add_paragraph(doc, f'{sub_num}) {sec["title"]}',
                           size=12, bold=True, color='2E5990',
                           indent_left=0.0,
                           before=10, after=6)
        else:
            # 3단계 — 더 작은 제목
            _add_paragraph(doc, f'· {sec["title"]}',
                           size=11, bold=True,
                           indent_left=0.8,
                           before=8, after=4)

        # 블록 렌더링
        blocks = sorted(sec.get('blocks', []),
                        key=lambda b: (b.get('display_order', 0), b.get('id', 0)))
        for b in blocks:
            _render_block(doc, b, depth)

        # 자식 섹션 (같은 페이지에 이어서)
        if sec.get('children'):
            _render_sections(doc, sec['children'], depth + 1, num)


def _add_horizontal_line(doc, color='808080', sz=6):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _render_block(doc, block, depth):
    """블록 타입별 렌더링"""
    bt = block.get('block_type')
    content = block.get('content') or {}
    # content가 문자열로 저장된 경우 (DB raw row) 디시리얼라이즈
    if isinstance(content, str):
        try:
            content = json.loads(content)
        except Exception:
            content = {}

    # depth에 따른 좌측 들여쓰기 (1단계 본문은 0.5cm, 2단계는 1cm)
    base_indent = 0.5 + 0.5 * depth

    if bt == 'paragraph':
        text = (content.get('text') or '').strip()
        if not text:
            return
        for line in text.split('\n'):
            _add_paragraph(doc, line, size=10.5,
                           indent_left=base_indent, after=4)

    elif bt == 'bullet_list':
        items = content.get('items') or []
        marker = content.get('marker') or 'bullet'
        # 레벨별 카운터
        counters = [0, 0, 0, 0]
        for it in items:
            if isinstance(it, str):
                text, indent = it, 0
            else:
                text = it.get('text', '')
                indent = max(0, min(3, it.get('indent', 0)))
            if not text.strip():
                continue
            counters[indent] += 1
            for k in range(indent + 1, 4):
                counters[k] = 0

            if marker == 'dash':       mk = '–'
            elif marker == 'number':   mk = f'{counters[indent]})'
            elif marker == 'alpha':    mk = f'{chr(96 + (counters[indent] - 1) % 26 + 1)})'
            else:                       mk = '•'

            line = f'{mk}  {text}'
            _add_paragraph(doc, line, size=10.5,
                           indent_left=base_indent + 0.6 * indent,
                           after=2)

    elif bt == 'table':
        _render_table_block(doc, content, base_indent)

    elif bt == 'image':
        _render_image_block(doc, content, base_indent)


def _render_table_block(doc, content, base_indent):
    headers = content.get('headers') or []
    rows = content.get('rows') or []
    if not headers and not rows:
        return

    n_cols = max(len(headers), max((len(r) for r in rows), default=0))
    if n_cols == 0:
        return

    # 헤더/행 길이 맞추기
    headers = (list(headers) + [''] * n_cols)[:n_cols]
    rows = [(list(r) + [''] * n_cols)[:n_cols] for r in rows]

    # 컬럼 너비 — 사용자 지정 있으면 비율로 적용, 없으면 균등
    col_widths_px = content.get('col_widths') or []
    total_cm = 16.0  # 본문 가용 폭 (A4 - 양쪽 여백)
    if col_widths_px and len(col_widths_px) == n_cols and sum(col_widths_px) > 0:
        total = sum(col_widths_px)
        col_cm = [total_cm * (w / total) for w in col_widths_px]
    else:
        col_cm = [total_cm / n_cols] * n_cols

    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 헤더
    for ci, h in enumerate(headers):
        cell = tbl.rows[0].cells[ci]
        cell.width = Cm(col_cm[ci])
        _set_cell_shading(cell, 'D9E2EC')
        _set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        _set_font(r, size=10, bold=True)

    # 데이터
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            cell.width = Cm(col_cm[ci])
            _set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # 표 셀에는 줄바꿈을 \n 으로 처리
            text = str(val or '')
            lines = text.split('\n')
            for li, ln in enumerate(lines):
                if li > 0:
                    p = cell.add_paragraph()
                r = p.add_run(ln)
                _set_font(r, size=10)

    # 표 위/아래 여백
    _add_paragraph(doc, '', before=2, after=4)


def _render_image_block(doc, content, base_indent):
    images = content.get('images') or []
    columns = max(1, min(4, int(content.get('columns', 2) or 2)))

    if not images:
        return

    # 이미지 그리드: N열 표로 배치
    n = len(images)
    n_rows = (n + columns - 1) // columns

    # 각 셀 너비: A4 본문폭 / columns
    total_cm = 16.0
    cell_cm = (total_cm - 0.3 * (columns - 1)) / columns

    # 캡션 행 포함 → 행 수 × 2 (이미지 행 + 캡션 행)
    tbl = doc.add_table(rows=n_rows * 2, cols=columns)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for idx, img in enumerate(images):
        ri = (idx // columns) * 2
        ci = idx % columns
        img_cell = tbl.rows[ri].cells[ci]
        cap_cell = tbl.rows[ri + 1].cells[ci]
        img_cell.width = Cm(cell_cm)
        cap_cell.width = Cm(cell_cm)

        # 이미지 삽입
        img_path = _resolve_image_path(img.get('url') or '', img.get('filename') or '')
        if img_path and os.path.exists(img_path):
            try:
                p = img_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Cm(cell_cm - 0.4))
            except Exception as e:
                p = img_cell.paragraphs[0]
                r = p.add_run(f'[이미지 로드 실패: {e}]')
                _set_font(r, size=9, color='B91C1C')
        else:
            p = img_cell.paragraphs[0]
            r = p.add_run('[이미지 없음]')
            _set_font(r, size=9, color='9CA3AF')

        # 캡션
        cap_p = cap_cell.paragraphs[0]
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = img.get('caption') or ''
        if caption:
            cap_run = cap_p.add_run(caption)
            _set_font(cap_run, size=9, color='4B5563')
            cap_run.italic = True

    # 빈 셀(마지막 행의 잉여) 처리는 자동
    _add_paragraph(doc, '', before=2, after=6)


def _resolve_image_path(url, filename):
    """저장된 이미지 URL → 파일 시스템 경로"""
    # url 예시: /static/uploads/dock/dock-1-1234-abc.jpg
    if url and url.startswith('/static/'):
        rel = url[len('/static/'):]
        # static 디렉터리 위치 — app 모듈에서 가져옴
        from app import app
        return os.path.join(app.static_folder, rel)
    # 직접 filename으로 fallback
    if filename:
        from app import app
        return os.path.join(app.static_folder, 'uploads', 'dock', filename)
    return None


# ─────────────────────────────────────────────────────────────
#  공개 함수 — 보고서 데이터 받아 docx 바이트 반환
# ─────────────────────────────────────────────────────────────
def build_docx(report: dict) -> bytes:
    """
    report: GET /api/dock-reports/<id> 응답과 동일한 구조
      · 메타: title, vessel_name, dock_no, shipyard, period_*, imo_no, gt, dwt,
              approval_*
      · sections: 평면 리스트, 각 항목에 blocks 포함
    """
    doc = Document()

    # 페이지 설정 (A4, 여백 2cm)
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width  = Mm(210)
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # 기본 폰트
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    rFonts.set(qn('w:ascii'), 'Malgun Gothic')
    rFonts.set(qn('w:hAnsi'), 'Malgun Gothic')

    # 페이지 번호 (푸터 우측)
    _add_page_number_footer(doc)

    # 섹션 트리 빌드
    sections_flat = report.get('sections') or []
    tree = _build_tree(sections_flat)

    # ① 표지
    _build_cover(doc, report)
    # ② 목차
    if tree:
        _build_toc(doc, tree)
    # ③ 본문
    _render_sections(doc, tree)

    # 바이트로 직렬화
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


def _build_tree(sections_flat):
    """평면 리스트 → parent_id 기반 트리"""
    by_id = {s['id']: dict(s, children=[]) for s in sections_flat}
    roots = []
    for s in by_id.values():
        if s.get('parent_id') and s['parent_id'] in by_id:
            by_id[s['parent_id']]['children'].append(s)
        else:
            roots.append(s)

    def sort_rec(lst):
        lst.sort(key=lambda x: (x.get('display_order', 0), x.get('id', 0)))
        for x in lst:
            sort_rec(x['children'])
    sort_rec(roots)
    return roots


def _add_page_number_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run()
        # 페이지 번호 필드
        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')
        run._element.append(fld_begin)

        instr = OxmlElement('w:instrText')
        instr.text = 'PAGE'
        run._element.append(instr)

        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')
        run._element.append(fld_end)

        _set_font(run, size=9, color='6B7280')
