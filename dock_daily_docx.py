"""Superintendent's Daily DD report(.docx) 읽기 — 순수 파서 + 일자 판정.

무엇을 하는가
-------------
조선소 감독이 매일 보내는 `Superintendets_Daily_DD_report_<날짜>.docx` 를 읽어
입거 Daily Report 카드로 넣을 **행 목록**을 만든다.  Flask·DB·번역·네트워크를
전혀 모르는 leaf 모듈이라 표 파싱과 날짜 판정만 단위테스트로 잠글 수 있다.

왜 문서의 날짜를 안 쓰는가 (형 지시 2026-08-23)
----------------------------------------------
라이브 문서의 `Reporting date-S/N:` 칸은 **비어 있다**(실측).  본문 날짜는 형식이
섞여 있고 오타도 실재한다 — `20/08.2026`(구분자 혼재), 초기미팅 줄에 전년도가
그대로 남은 `19/12/2025`.  그래서 "문서에서 날짜를 뽑아 보고서를 찾는" 방향은
**틀린 보고서를 덮는 사고**로 끝난다.  구조를 뒤집었다: 사람이 그 날짜의 빈
보고서를 먼저 만들고, 그 보고서가 기준일 `D` 를 준다.  파서는 문서에서 날짜를
**읽되 절대 결정하지 않는다.**

🔴 `Date finish` 는 완료 기록이 아니다 (실측)
--------------------------------------------
8/20 라이브 문서에서 날짜가 있는 행 **전부** `Date finish` 가 `Scheduled finish
date` 와 글자까지 같다 — `01.09/01.09`, `21.08/21.08`, `28.08/28.08`,
`25.08/25.08`.  즉 감독은 계획일을 두 칸에 같이 적고 있고, `finish` 를 완료로
읽으면 아직 안 끝난 작업에 **"금일 완료"** 를 찍어 형이 그대로 사내 메일로
보낸다.  그래서 두 칸이 **다를 때만** `finish` 를 실제 완료로 신뢰한다
(`actual`).  같으면 계획(`plan`)일 뿐이다.

이 템플릿이 이렇게 채워지는 동안 "금일 완료" 표시는 거의 안 나온다.  그게 맞다 —
문서에 완료 사실이 안 적혀 있으니 완료는 형이 손으로 쓰는 게 정직하다.

일자 판정 (이 템플릿은 행이 매일 누적되므로 이게 곧 중복 방지다)
--------------------------------------------------------------
`start` / `plan` / `actual` 로 정한다 (`D` = 보고서 일자).

  · `actual == D`                             → 포함 (당일 완료)
  · `start == D`                              → 포함 (당일 착수)
  · `start > D`                               → 제외 (미래)
  · `start < D`, `actual` 있고 `< D`           → 제외 (이미 끝남)
  · `start < D`, `actual` 있고 `> D`           → 포함 (완료 예정일이 아직 안 옴)
  · `start < D`, `actual` 없고 `plan >= D`     → 포함 (진행 중)
  · `start < D`, `actual` 없고 `plan < D`      → 제외 (계획일 지남, 기록 없음)
  · `start < D`, `actual`·`plan` 없음          → 제외 (단발 과거).
       단 본문에 진행 문구(in progress/ongoing/…)가 있으면 포함
  · 날짜를 하나도 못 읽음                       → `unknown`.  **버리지 않는다.**

🔴 `plan`·`actual` 빈칸을 곧바로 "진행 중"으로 보면 안 된다.  실측 8/20 문서의
19.08 행 4개(도착·입국수속·초기미팅·자재적재)가 전부 빈칸이고, 그걸 진행 중으로
읽으면 형이 지운 전날 내용이 매일 되살아난다.  계획완료일이 있는 행만 다일
작업으로 신뢰한다.

🔴 `unknown` 을 조용히 버리지 않는 이유: 실측에 `UTM Service` 처럼 3칸이 전부 빈
행이 있다.  버리면 형은 문서에 있던 작업이 사라진 걸 모르고 메일을 보낸다.
"""
import datetime as dt
import hashlib
import re

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

#: 표 **위의 제목 문단** → (섹션 key, 화면 라벨).  문서 템플릿의 고정 문자열이다.
#: 정규식으로 보는 이유는 감독이 대소문자·공백·`3rd`/`third` 를 흔들기 때문이다.
HEADING_RULES = (
    (re.compile(r'leading\s+deck\s+works.*yard', re.I), 'shipyard', 'Deck (Yard)'),
    (re.compile(r'leading\s+engine\s+works.*yard', re.I), 'shipyard', 'Engine (Yard)'),
    # 🔴 `works` 앞의 `deck`/`engine` 만 관용한다.  같은 보고서의 PDF 템플릿(TDF-04.7c)은
    # `Leading Engine Works done by the 3rd party` / `... the Crew` 라고 쓴다(실측).
    # 붙여 읽던 옛 규칙에서는 이 두 표가 통째로 비고 섹션으로 흘렀다.
    # 🔴 임의의 한 단어(`\w+`)를 받으면 안 된다(올마이트 지적 2026-08-23) -- 기존 `.docx`
    # 의 엉뚱한 제목까지 vendor/crew 로 새로 분류돼 이미 비고로 들어간 카드가 다른
    # 섹션에 중복된다.  라벨은 그대로다 -- 바꾸면 `row_key` 가 재키잉돼 중복된다.
    (re.compile(r'leading\s+(?:deck\s+|engine\s+)?works.*(3rd|third)\s*part', re.I),
     'vendor', '3rd party'),
    (re.compile(r'leading\s+(?:deck\s+|engine\s+)?works.*crew', re.I), 'crew', 'Crew'),
)

#: 제목처럼 보이지만 규칙에 안 맞을 때 쓰는 신호.  이 단어들이 있으면 작업표 제목
#: 후보로 보고 `unmapped` 로 올린다(조용히 무시하지 않는다).
HEADING_HINT = re.compile(r'leading\s+works|works\s+done\s+by', re.I)

#: 진행 중임을 본문이 직접 말하는 경우.  날짜칸이 비어도 이건 신뢰한다.
PROGRESS_RE = re.compile(r'in\s*progress|ongoing|continu|under\s*way|underway|to\s+be\s+cont', re.I)

#: 구분자 혼재(`20/08.2026`)와 2자리 연도를 관용한다.  형식을 못 읽으면 `BAD`.
_DATE_RE = re.compile(r'(\d{1,2})\s*[./\-]\s*(\d{1,2})\s*[./\-]\s*(\d{2,4})')

BAD_DATE = 'BAD'

COL_ALIASES = (
    # `sn` 은 `.docx` 경로에선 안 쓰지만(번호칸을 읽지 않는다) PDF 경로가 "번호도
    # 날짜도 없고 설명만 있는 행" = 페이지 넘김에 잘린 조각인지 가리는 데 쓴다.
    ('sn', re.compile(r'^\s*s\s*/?\s*n\b|^\s*no\.?\s*$', re.I)),
    ('desc', re.compile(r'descrip', re.I)),
    ('start', re.compile(r'date\s*start|start\s*date', re.I)),
    ('sched', re.compile(r'schedul', re.I)),
    ('finish', re.compile(r'date\s*finish|finish\s*date', re.I)),
)


def tolerant_date(text):
    """`None`(빈칸) · `BAD_DATE`(못 읽음) · `datetime.date` 중 하나.

    🔴 `BAD_DATE` 와 `None` 을 합치면 안 된다.  빈칸은 계약상 의미가 있고(진행 중
    판정에 쓰인다) 오타는 사람이 봐야 하는 것이다.
    """
    s = (text or '').strip()
    if not s:
        return None
    m = _DATE_RE.search(s)
    if not m:
        return BAD_DATE
    day, month, year = (int(x) for x in m.groups())
    if year < 100:
        year += 2000
    try:
        return dt.date(year, month, day)
    except ValueError:
        return BAD_DATE


def row_key(scope, description):
    """같은 작업을 재업로드에서도 같은 것으로 알아보는 안정 키.

    행 번호를 쓰지 않는다 — 감독이 행을 지우거나 끼워 넣으면 번호가 전부 밀려서
    같은 작업이 새 카드로 또 들어온다.

    🔴 `scope` 는 섹션 key 가 아니라 **표 라벨**이다(올마이트 지적 2026-08-23).
    Deck 표와 Engine 표는 둘 다 `shipyard` 로 가므로, 섹션 key 로 묶으면 두 표에
    같은 문장이 있을 때 키가 겹쳐 뒤 행이 앞 행 카드를 덮고 **한 건이 조용히
    사라진다.**  라벨(`Deck (Yard)`/`Engine (Yard)`)은 규칙표에서 오는 고정값이라
    감독이 제목 문구를 흔들어도 키가 안 변한다.
    """
    norm = re.sub(r'\s+', ' ', (description or '')).strip().lower()
    raw = '%s|%s' % (re.sub(r'\s+', ' ', (scope or '')).strip().lower(), norm)
    return hashlib.sha256(raw.encode('utf-8')).hexdigest()[:12]


def _iter_blocks(doc):
    """문단과 표를 **문서 순서대로**.  제목→표 짝짓기가 순서에 달려 있다."""
    for child in doc.element.body.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, doc)
        elif child.tag == qn('w:tbl'):
            yield Table(child, doc)


def heading_target(text):
    """제목 문단 → `(섹션 key, 라벨)`.  PDF 경로도 **이 규칙표를 그대로** 쓴다."""
    for pattern, key, label in HEADING_RULES:
        if pattern.search(text):
            return key, label
    return None, None


def _columns(header_cells):
    """헤더 텍스트로 컬럼 위치를 찾는다.

    위치를 못박지 않는 이유: 번호 컬럼이 있는 표와 없는 표가 같은 문서에 섞여
    있다(실측).  위치로 세면 3rd party 표에서 Description 을 번호칸으로 읽는다.
    """
    found = {}
    for idx, cell in enumerate(header_cells):
        for name, pattern in COL_ALIASES:
            if name not in found and pattern.search(cell):
                found[name] = idx
    return found if 'desc' in found else None


def parse(path):
    """`{'groups': [...], 'unmapped_headings': [...]}`.

    각 group = `{'heading', 'label', 'section_key', 'rows': [...]}`,
    각 row = `{'row_key', 'no', 'desc', 'start', 'sched', 'finish'}` (문자열 원본).
    판정은 `classify()` 가 따로 한다 — 파싱과 판정을 섞으면 기준일을 바꿔 가며
    테스트할 수 없다.
    """
    doc = Document(path)
    groups, unmapped = [], []
    pending = None  # (heading, label, section_key) — 직전에 본 제목
    for block in _iter_blocks(doc):
        if isinstance(block, Paragraph):
            text = re.sub(r'\s+', ' ', block.text).strip()
            if not text:
                continue
            key, label = heading_target(text)
            if key:
                pending = (text, label, key)
            elif HEADING_HINT.search(text):
                pending = (text, text, None)
                unmapped.append(text)
            continue
        if not pending:
            continue
        rows = [[c.text.strip() for c in r.cells] for r in block.rows]
        if not rows:
            continue
        cols = _columns(rows[0])
        if not cols:
            # 제목 바로 뒤인데 Description 헤더가 없는 표는 작업표가 아니다
            # (날씨·비용 표). 제목을 소비해 다음 표로 새지 않게 한다.
            pending = None
            continue
        heading, label, section_key = pending
        pending = None
        out_rows = []
        for cells in rows[1:]:
            def cell(name):
                idx = cols.get(name)
                return cells[idx] if idx is not None and idx < len(cells) else ''
            desc = cell('desc')
            if not desc.strip():
                continue
            out_rows.append({
                'row_key': row_key(label or heading, desc),
                'desc': re.sub(r'\s+', ' ', desc).strip(),
                'start': cell('start'), 'sched': cell('sched'), 'finish': cell('finish'),
            })
        if out_rows:
            groups.append({'heading': heading, 'label': label,
                           'section_key': section_key, 'rows': out_rows})
    return {'groups': groups, 'unmapped_headings': unmapped}


def _dates(row):
    """`(start, plan, actual, bad)`.

    `actual` 은 `Date finish` 가 `Scheduled finish date` 와 **다를 때만** 채운다.
    같으면 계획을 두 칸에 복사한 것이라 완료 근거가 못 된다(모듈 docstring 실측).
    """
    start = tolerant_date(row.get('start'))
    sched = tolerant_date(row.get('sched'))
    finish = tolerant_date(row.get('finish'))
    bad = BAD_DATE in (start, sched, finish)
    if bad:
        return None, None, None, True
    actual = finish if (finish is not None and finish != sched) else None
    return start, sched, actual, False


def classify(row, report_date):
    """`('include'|'exclude'|'unknown', 사유코드)`."""
    start, plan, actual, bad = _dates(row)
    if bad:
        return 'unknown', 'date_unreadable'
    if actual == report_date:
        return 'include', 'finished_today'
    if start == report_date:
        return 'include', 'started_today'
    if start is None:
        return 'unknown', 'no_start_date'
    if start > report_date:
        return 'exclude', 'future'
    # start < report_date
    if actual is not None:
        # 🔴 완료일이 기준일보다 **뒤**면 그 작업은 이 날 아직 돌고 있다(올마이트 지적
        # 2026-08-23).  옛 코드는 `actual` 이 있다는 이유만으로 "이미 끝남" 으로 빼서,
        # 진행 중인 공사가 보고서에서 통째로 누락됐다.
        if actual > report_date:
            return 'include', 'finish_in_future'
        return 'exclude', 'finished_earlier'
    if plan is not None:
        if plan >= report_date:
            return 'include', 'in_progress_scheduled'
        return 'exclude', 'plan_date_passed'
    if PROGRESS_RE.search(row.get('desc') or ''):
        return 'include', 'in_progress_text'
    return 'exclude', 'past_one_off'


def marker(row, report_date):
    """카드 문장 끝에 붙일 상태 표시.  없으면 빈 문자열.

    🔴 계획일만 있는 행에 "완료" 라고 쓰지 않는다.  형이 이 문장을 그대로 사내
    메일로 보내기 때문에 문서에 없는 사실을 만들면 안 된다.
    """
    start, plan, actual, bad = _dates(row)
    if bad:
        return ''
    if actual == report_date:
        return '금일 완료'
    # 완료칸이 기준일보다 뒤면 그건 예정이다. `계획 완료`(Scheduled) 와 구분해서 쓴다 --
    # 두 칸의 날짜가 다를 때 어느 쪽을 적은 건지 형이 알아야 한다.
    if isinstance(actual, dt.date) and actual > report_date:
        return '진행 중, 완료 예정 %02d.%02d' % (actual.month, actual.day)
    plan_txt = '계획 완료 %02d.%02d' % (plan.month, plan.day) if isinstance(plan, dt.date) else ''
    if start == report_date:
        return plan_txt if plan != start else ''
    if isinstance(start, dt.date) and start < report_date:
        return '진행 중, %s' % plan_txt if plan_txt else '진행 중'
    return plan_txt


def scan(path, report_date):
    """파싱 + 판정을 합친 미리보기 페이로드.

    `report_date` 는 `datetime.date`.  결과의 모든 행은 `verdict`/`reason`/
    `marker` 를 갖고, 제외된 행도 **버리지 않고** 그대로 실려 온다 — 형이 화면에서
    "왜 안 들어왔는지" 를 보고 필요하면 직접 고를 수 있어야 한다.
    """
    return judge(parse(path), report_date)


def judge(parsed, report_date):
    """파싱 결과에 판정을 얹는다(제자리 수정 후 그대로 반환).

    `.pdf` 경로(`dock_daily_pdf.scan`)가 **이 함수를 그대로** 부른다 — 판정 규칙이
    파일형식에 따라 갈리면 같은 보고서가 형식만 바꿔도 다르게 들어온다.
    """
    counts = {'include': 0, 'exclude': 0, 'unknown': 0}
    for group in parsed['groups']:
        for row in group['rows']:
            verdict, reason = classify(row, report_date)
            row['verdict'], row['reason'] = verdict, reason
            row['marker'] = marker(row, report_date) if verdict == 'include' else ''
            counts[verdict] += 1
    parsed['counts'] = counts
    parsed['report_date'] = report_date.isoformat()
    return parsed


# ─────────────────────────────────────────────────────────────────────
#  문서에 박힌 사진 (형 지시 2026-08-23 "사진도 자동으로")
# ─────────────────────────────────────────────────────────────────────
#: 캡션 칸의 **템플릿 자리표시자**.  라이브 문서 끝에는 채우지 않은
#: `Photo 1: Description` … `Photo 16: Description` 이 16줄 그대로 남아 있다(실측).
#: 이걸 캡션으로 쓰면 형이 사내 메일로 `Photo 7: Description` 을 보낸다.
PHOTO_PLACEHOLDER_RE = re.compile(r'^\s*photo\s*\d*\s*[:.]?\s*(description)?\s*$', re.I)

#: 한 문서에서 받아올 사진 수 상한.  넘으면 **자르되 몇 장을 잘랐는지 올린다**
#: (`photo_limit`) -- 조용히 자르면 형은 문서에 있던 사진이 빠진 걸 모른다.
PHOTO_MAX = 40

#: 사진으로 인정할 magic.  raw bitmap(Flate) 은 재인코딩이 필요한데 실측 문서에는
#: 없다 -- 검증 못 한 경로를 만드는 대신 `skipped` 로 세어 화면에 올린다.
PHOTO_MAGIC = ((b'\xff\xd8\xff', 'jpg', 'image/jpeg'),
               (b'\x89PNG\r\n\x1a\n', 'png', 'image/png'))


def photo_kind(data):
    """`(확장자, mime)` 또는 `(None, None)`."""
    for magic, ext, mime in PHOTO_MAGIC:
        if data[:len(magic)] == magic:
            return ext, mime
    return None, None


def photo_caption(text):
    """캡션 후보 문자열 → 쓸 만한 캡션.  자리표시자·빈칸은 `''`."""
    clean = re.sub(r'\s+', ' ', text or '').strip()
    return '' if PHOTO_PLACEHOLDER_RE.match(clean) else clean


def photo_key(digest, used):
    """짧은 표시·전송용 키.  접기 판정은 **full sha256** 으로 하고 이것만 자른다.

    🔴 문서 안에서 **유일**해야 한다.  두 사진이 같은 키를 받으면 형이 한 장을 골라도
       서버가 두 장을 넣는다(`photo_keys` 는 키로 고른다).  앞자리가 겹치면 길이를
       늘리고, 끝까지 겹치면 full digest 를 쓴다.
    """
    for size in range(12, 65, 8):
        short = digest[:size]
        if short not in used:
            return short
    return digest


def _cell_rids(cell):
    """이 칸 **자기 문단**의 그림만.

    🔴 `iter()` 로 훑으면 중첩 표 안 그림까지 딸려 온다.  `_all_tables` 가 중첩 표를
       따로 한 번 더 돌기 때문에 같은 사진이 바깥칸·안쪽칸에서 **두 번** 발견되고,
       바이트로 접히기는 하지만 `duplicates` 가 부풀고 캡션은 먼저 만난 **바깥칸**
       것이 채택된다 -- 엉뚱한 문장이 사진에 붙어 그대로 조선소로 나간다.
    """
    return [b.get(qn('r:embed')) for b in cell._tc.xpath('./w:p//a:blip')
            if b.get(qn('r:embed'))]


def _all_tables(doc):
    """중첩 표까지 문서 순서대로.  `doc.tables` 는 최상위만 준다."""
    out = []

    def walk(tables):
        for table in tables:
            out.append(table)
            for row in table.rows:
                for cell in row.cells:
                    walk(cell.tables)

    walk(doc.tables)
    return out


def _grid(table):
    """`[[{'text','rids'}]]`.  병합 칸은 `row.cells` 가 반복해 주므로 그대로 쓴다."""
    grid = []
    for row in table.rows:
        try:
            cells = list(row.cells)
        except (IndexError, ValueError):       # pragma: no cover - 손상 표
            continue
        grid.append([{'text': c.text, 'rids': _cell_rids(c)} for c in cells])
    return grid


def _caption_for(grid, i, j):
    """`(i,j)` 칸에 있는 사진의 캡션.

    실측 템플릿이 두 가지다 -- 둘 다 맞춰야 한다.
      · `Superintendets_Daily_DD_report_20.08` : `[캡션][사진]` **같은 행**.
      · `TDF-04.7c … 21.08`                    : `[사진][사진]` 다음 행이 `[캡션][캡션]`.
    🔴 없는 캡션을 만들지 않는다.  후보가 없으면 빈 캡션으로 사진만 들여온다 --
       엉뚱한 문장을 붙이면 형이 그 캡션을 그대로 조선소에 보낸다.
    """
    own = photo_caption(grid[i][j]['text'])
    if own:
        return own
    row = grid[i]
    # 같은 행: 가까운 칸부터.  사진이 든 칸도 후보다 -- 라이브 문서에 캡션칸으로
    # 사진 한 장이 더 붙은 행이 있고(20.08 3행), 그 문장은 옆 사진의 캡션이기도 하다.
    for dist in range(1, len(row)):
        for k in (j - dist, j + dist):
            if 0 <= k < len(row) and k != j:
                text = photo_caption(row[k]['text'])
                if text:
                    return text
    # 아래 행 같은 열: 사진행/캡션행이 번갈아 오는 템플릿.
    if i + 1 < len(grid) and j < len(grid[i + 1]):
        below = grid[i + 1][j]
        if not below['rids']:
            return photo_caption(below['text'])
    return ''


def photos(path):
    """문서에 박힌 사진.  `{'photos': [...], 'captions': True, ...}`.

    각 사진 = `{'photo_key', 'caption', 'data', 'ext', 'mime', 'size'}`.

    🔴 **표 안에 있는 사진만** 가져온다(실측 근거).  라이브 20.08 문서는 본문
       사진 11장 중 10장이 사진표 안에 있고 1장은 표 밖 머리글 그림이며, 21.08
       문서의 `image12.png` 는 body 에 아예 없는 머리글(header part) 로고다.
       "미디어 폴더를 전부 가져오기" 로 짜면 형의 보고서에 로고가 사진으로 박힌다.
    🔴 동일 바이트는 한 장으로 접는다(`duplicates`).  같은 사진이 두 칸에 붙어
       있으면 첨부도 두 개가 되고, 형은 왜 같은 사진이 두 장인지 알 수 없다.
    """
    doc = Document(path)
    part = doc.part
    out, seen, keys, dupes, skipped, over = [], set(), set(), 0, 0, 0
    for table in _all_tables(doc):
        grid = _grid(table)
        for i, row in enumerate(grid):
            for j, cell in enumerate(row):
                for rid in cell['rids']:
                    image = part.related_parts.get(rid)
                    if image is None:          # pragma: no cover - 깨진 관계
                        skipped += 1
                        continue
                    data = image.blob
                    ext, mime = photo_kind(data or b'')
                    if not ext:
                        skipped += 1
                        continue
                    # 🔴 접기 판정은 **full digest** 다.  12자로 접으면 앞자리가 겹친
                    #    서로 다른 사진 한 장이 조용히 사라진다.
                    digest = hashlib.sha256(data).hexdigest()
                    if digest in seen:
                        dupes += 1
                        continue
                    seen.add(digest)
                    if len(out) >= PHOTO_MAX:
                        over += 1
                        continue
                    key = photo_key(digest, keys)
                    keys.add(key)
                    out.append({'photo_key': key, 'caption': _caption_for(grid, i, j),
                                'data': data, 'ext': ext, 'mime': mime, 'size': len(data)})
    return {'photos': out, 'captions': True, 'duplicates': dupes,
            'skipped': skipped, 'photo_limit': over, 'letterhead': 0}
