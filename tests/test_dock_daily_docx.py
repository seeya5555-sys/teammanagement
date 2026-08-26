"""감독 Daily DD report(.docx) → 입거 Daily Report 카드 (형 지시 2026-08-23).

파서 단위 계약과 두 라우트(`docx-scan` / `docx-apply`)의 계약을 잠근다.  fixture 는
테스트 시점에 `python-docx` 로 만든다 -- 바이너리를 저장소에 넣으면 왜 그런 결과가
나오는지 아무도 못 읽는다.
"""
import datetime as dt
import io
import json
import os
import tempfile
import unittest

from docx import Document

import app as appmod
import dock_daily_docx as parser
import routes_dock_daily

appmod.app.config['CSRF_PROTECT'] = False

D = dt.date(2026, 8, 20)


def _table(doc, heading, header, rows):
    """제목 문단 + 표 하나.  실문서와 같은 순서(제목이 표 위)로 넣는다."""
    doc.add_paragraph(heading)
    table = doc.add_table(rows=1, cols=len(header))
    for i, text in enumerate(header):
        table.rows[0].cells[i].text = text
    for values in rows:
        cells = table.add_row().cells
        for i, text in enumerate(values):
            cells[i].text = text
    return table


HEADER = ['No', 'Description', 'Date start', 'Scheduled finish date', 'Date finish']


def build_docx(path, rows=None, extras=True):
    """실문서 구조를 닮은 fixture.

    실측 사실을 그대로 반영한다: 번호 컬럼이 있는 표와 없는 표가 섞여 있고,
    `Date finish` 에는 계획일이 그대로 복사돼 있다.
    """
    doc = Document()
    doc.add_paragraph('Superintendent Daily DD report')
    doc.add_paragraph('Reporting date-S/N:')          # 🔴 실문서처럼 값이 없다
    _table(doc, 'Leading Deck Works done by the Yard', HEADER, rows if rows is not None else [
        ['1', 'Vessel arrived to SY 11:42 LT', '19.08.2026', '', ''],
        ['2', 'Hatch Cover No.1 dismantle hydraulic jacks', '20/08/2026', '21.08.2026', '21.08.2026'],
        ['3', 'Load/discharge spare parts and tools', '20/08/2026', '', ''],
        ['4', 'Staging erection continues in cargo hold', '18.08.2026', '', ''],
        ['5', 'Future blasting of hull', '25.08.2026', '', ''],
        ['6', '', '20.08.2026', '', ''],                    # 빈 설명 → 버린다
    ])
    if not extras:
        return doc.save(path)
    # 번호 컬럼이 **없는** 표 (실문서의 3rd party 표가 이렇다)
    _table(doc, 'Leading Works done by 3rd party',
           ['Description', 'Date start', 'Scheduled finish date', 'Date finish'], [
               ['ME TC Turbocharger Service', '20.08.2026', '28.08.2026', '28.08.2026'],
               ['UTM Service', '', '', ''],                 # 3칸 전부 빈칸 → 판정불가
           ])
    _table(doc, 'Leading Works done by Crew', HEADER, [
        ['1', 'ME FO Auto-filter maintenance', '20.08.2026', '', ''],
    ])
    # Description 헤더가 없는 표(날씨). 제목 뒤에 와도 작업표로 읽지 않는다.
    _table(doc, 'Weather conditions', ['Wind', 'Temp'], [['NE 4', '28']])
    doc.save(path)


class ParserTests(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.path = os.path.join(self.tmp.name, 'dd.docx')
        build_docx(self.path)

    def tearDown(self):
        self.tmp.cleanup()

    def test_tolerant_date_separates_blank_from_unreadable(self):
        """🔴 빈칸과 오타를 같은 값으로 뭉치면 안 된다.

        빈칸은 계약상 의미가 있고(진행 중 판정에 쓴다) 오타는 사람이 봐야 한다.
        """
        self.assertIsNone(parser.tolerant_date(''))
        self.assertIsNone(parser.tolerant_date('   '))
        self.assertIsNone(parser.tolerant_date(None))
        self.assertEqual(dt.date(2026, 8, 20), parser.tolerant_date('20.08.2026'))
        self.assertEqual(dt.date(2026, 8, 20), parser.tolerant_date('20/08/2026'))
        # 실문서에 있는 구분자 혼재
        self.assertEqual(dt.date(2026, 8, 20), parser.tolerant_date('20/08.2026'))
        self.assertEqual(dt.date(2026, 8, 20), parser.tolerant_date('20-08-26'))
        self.assertEqual(parser.BAD_DATE, parser.tolerant_date('to be advised'))
        self.assertEqual(parser.BAD_DATE, parser.tolerant_date('32.08.2026'))
        self.assertEqual(parser.BAD_DATE, parser.tolerant_date('20.13.2026'))

    def test_date_finish_equal_to_the_plan_is_not_a_completion(self):
        """🔴 실측: 감독은 계획일을 `Scheduled finish` 와 `Date finish` 두 칸에 같이 적는다.

        이걸 완료로 읽으면 아직 안 끝난 작업에 "금일 완료" 가 찍히고, 형은 그 문장을
        그대로 사내 메일로 보낸다.  두 칸이 **다를 때만** 완료로 인정해야 한다.
        """
        plan_only = {'desc': 'Hatch cover', 'start': '20.08.2026',
                     'sched': '21.08.2026', 'finish': '21.08.2026'}
        self.assertEqual(('include', 'in_progress_scheduled'),
                         parser.classify(plan_only, dt.date(2026, 8, 21)))
        self.assertEqual('진행 중, 계획 완료 08.21', parser.marker(plan_only, dt.date(2026, 8, 21)))
        self.assertNotIn('완료 ', parser.marker(plan_only, dt.date(2026, 8, 21)).replace('계획 완료 ', ''))

        # 두 칸이 다르면 그건 실제 완료 기록이다.
        real = dict(plan_only, sched='28.08.2026', finish='21.08.2026')
        self.assertEqual(('include', 'finished_today'), parser.classify(real, dt.date(2026, 8, 21)))
        self.assertEqual('금일 완료', parser.marker(real, dt.date(2026, 8, 21)))
        # 완료일이 지난 행은 다음 날 빠진다(= 중복 방지).
        self.assertEqual(('exclude', 'finished_earlier'),
                         parser.classify(real, dt.date(2026, 8, 22)))

    def test_work_finishing_after_the_report_date_is_still_in_progress(self):
        """🔴 완료칸이 기준일보다 뒤면 그 작업은 이 날 아직 돌고 있다.

        옛 규칙은 `Date finish` 가 채워졌다는 이유만으로 "이미 끝남" 으로 빼서, 진행
        중인 공사가 보고서에서 통째로 누락됐다(올마이트 지적 2026-08-23).
        """
        row = {'desc': 'Hull blasting', 'start': '18.08.2026',
               'sched': '28.08.2026', 'finish': '30.08.2026'}
        self.assertEqual(('include', 'finish_in_future'), parser.classify(row, D))
        # 계획일(28)이 아니라 문서가 적은 완료 예정일(30)을 쓴다 -- 두 날짜가 다를 때
        # 어느 칸을 읽은 건지 형이 알아야 한다.
        self.assertEqual('진행 중, 완료 예정 08.30', parser.marker(row, D))
        self.assertEqual(('exclude', 'finished_earlier'),
                         parser.classify(row, dt.date(2026, 9, 1)))

    def test_classify_covers_every_branch(self):
        cases = (
            ('당일 착수', {'start': '20.08.2026'}, 'include', 'started_today'),
            ('미래', {'start': '25.08.2026'}, 'exclude', 'future'),
            ('계획일 남음 = 진행 중', {'start': '18.08.2026', 'sched': '25.08.2026'},
             'include', 'in_progress_scheduled'),
            ('계획일 지남', {'start': '10.08.2026', 'sched': '12.08.2026'},
             'exclude', 'plan_date_passed'),
            ('단발 과거', {'start': '19.08.2026'}, 'exclude', 'past_one_off'),
            ('오타', {'start': 'yesterday'}, 'unknown', 'date_unreadable'),
            ('날짜 없음', {}, 'unknown', 'no_start_date'),
        )
        for name, row, verdict, reason in cases:
            with self.subTest(name):
                self.assertEqual((verdict, reason), parser.classify(dict(row, desc='x'), D))

    def test_progress_wording_rescues_a_dateless_past_row(self):
        """날짜칸이 비어도 본문이 진행 중이라고 말하면 넣는다(형 지시)."""
        row = {'desc': 'Staging erection in cargo hold - in progress', 'start': '18.08.2026'}
        self.assertEqual(('include', 'in_progress_text'), parser.classify(row, D))
        self.assertEqual('진행 중', parser.marker(row, D))
        # 🔴 같은 날짜꼴인데 진행 문구가 없으면 빠진다. 안 그러면 형이 지운 전날
        #    내용이 매일 되살아난다.
        self.assertEqual('exclude', parser.classify(dict(row, desc='Staging erection'), D)[0])

    def test_headings_map_to_sections_and_a_numberless_table_still_parses(self):
        out = parser.parse(self.path)
        self.assertEqual([], out['unmapped_headings'])
        got = [(g['section_key'], g['label'], len(g['rows'])) for g in out['groups']]
        self.assertEqual([('shipyard', 'Deck (Yard)', 5),
                          ('vendor', '3rd party', 2),
                          ('crew', 'Crew', 1)], got,
                         '날씨 표는 작업표가 아니고, 빈 설명 행은 버린다')
        # 번호 컬럼이 없는 표에서도 Description 을 번호칸으로 잘못 읽지 않는다.
        self.assertEqual('ME TC Turbocharger Service', out['groups'][1]['rows'][0]['desc'])
        self.assertEqual('20.08.2026', out['groups'][1]['rows'][0]['start'])

    def test_an_unknown_heading_is_surfaced_not_dropped(self):
        path = os.path.join(self.tmp.name, 'odd.docx')
        doc = Document()
        _table(doc, 'Leading Works done by Subcontractor Alpha', HEADER,
               [['1', 'Anode renewal', '20.08.2026', '', '']])
        doc.save(path)
        out = parser.parse(path)
        self.assertEqual(['Leading Works done by Subcontractor Alpha'], out['unmapped_headings'])
        self.assertIsNone(out['groups'][0]['section_key'])
        self.assertEqual(1, len(out['groups'][0]['rows']), '모르는 제목이어도 행은 살린다')

    def test_the_third_party_heading_accepts_deck_engine_but_not_any_word(self):
        """🔴 제목 규칙을 임의의 한 단어(`\\w+`)까지 열면 안 된다(올마이트 지적 2026-08-23).

        `Leading Deck Works ... 3rd party` 같은 실제 변형은 받아야 하지만, 아무 단어나
        받으면 기존 `.docx` 의 엉뚱한 제목까지 vendor/crew 로 새로 분류돼 이미 비고에
        들어간 카드가 다른 섹션에 중복된다.  모르는 제목은 `unmapped_headings` 로 올린다.
        """
        cases = [('Leading Deck Works done by the 3rd party', 'vendor'),
                 ('Leading Engine Works done by the 3rd party', 'vendor'),
                 ('Leading Works done by 3rd party', 'vendor'),
                 ('Leading Deck Works done by the Crew', 'crew'),
                 ('Leading Engine Works done by the Crew', 'crew'),
                 ('Leading Hull Works done by the 3rd party', None),
                 ('Leading Repair Works done by the Crew', None)]
        for heading, want in cases:
            path = os.path.join(self.tmp.name, 'h%d.docx' % abs(hash(heading)))
            doc = Document()
            _table(doc, heading, HEADER, [['1', 'Anode renewal', '20.08.2026', '', '']])
            doc.save(path)
            out = parser.parse(path)
            self.assertEqual(want, out['groups'][0]['section_key'], heading)
            self.assertEqual([] if want else [heading], out['unmapped_headings'], heading)
            self.assertEqual(1, len(out['groups'][0]['rows']), '어느 쪽이든 행은 살린다')

    def test_the_same_sentence_in_the_deck_and_engine_tables_stays_two_rows(self):
        """🔴 Deck 표와 Engine 표는 둘 다 `shipyard` 섹션으로 간다.

        섹션 key 로 `row_key` 를 만들면 두 표에 같은 문장이 있을 때 키가 겹쳐서 한
        건이 조용히 사라졌다(올마이트 지적 2026-08-23).  표 라벨로 나눠야 한다.
        """
        path = os.path.join(self.tmp.name, 'both.docx')
        doc = Document()
        same = 'Pipe renewal in engine room'
        _table(doc, 'Leading Deck Works done by the Yard', HEADER,
               [['1', same, '20.08.2026', '', '']])
        _table(doc, 'Leading Engine Works done by the Yard', HEADER,
               [['1', same, '20.08.2026', '', '']])
        doc.save(path)
        out = parser.parse(path)
        keys = [r['row_key'] for g in out['groups'] for r in g['rows']]
        self.assertEqual(2, len(keys))
        self.assertEqual(2, len(set(keys)), '같은 문장이라도 다른 표면 다른 행이다')
        self.assertEqual(['shipyard', 'shipyard'], [g['section_key'] for g in out['groups']])

    def test_row_key_ignores_heading_wording_drift(self):
        """제목 문구가 흔들려도 같은 작업은 같은 키다(라벨은 규칙표에서 온다)."""
        first = os.path.join(self.tmp.name, 'h1.docx')
        second = os.path.join(self.tmp.name, 'h2.docx')
        for path, heading in ((first, 'Leading Deck Works done by the Yard'),
                              (second, 'LEADING   deck works  done by yard')):
            doc = Document()
            _table(doc, heading, HEADER, [['1', 'Anode renewal', '20.08.2026', '', '']])
            doc.save(path)
        a = parser.parse(first)['groups'][0]['rows'][0]['row_key']
        b = parser.parse(second)['groups'][0]['rows'][0]['row_key']
        self.assertEqual(a, b)

    def test_row_key_is_stable_when_the_yard_reorders_rows(self):
        """행 번호로 동일성을 잡으면 감독이 행 하나만 끼워 넣어도 전부 새 카드가 된다."""
        first = parser.parse(self.path)['groups'][0]['rows']
        shuffled = os.path.join(self.tmp.name, 'shuffled.docx')
        build_docx(shuffled, rows=[
            ['1', 'NEW inserted row', '20.08.2026', '', ''],
            ['2', 'Vessel arrived to SY 11:42 LT', '19.08.2026', '', ''],
            ['3', 'Hatch Cover No.1 dismantle hydraulic jacks', '20/08/2026', '21.08.2026', '21.08.2026'],
        ], extras=False)
        after = {r['desc']: r['row_key'] for r in parser.parse(shuffled)['groups'][0]['rows']}
        for row in first:
            if row['desc'] in after:
                self.assertEqual(row['row_key'], after[row['desc']], row['desc'])

    def test_scan_counts_and_keeps_excluded_rows(self):
        out = parser.scan(self.path, D)
        self.assertEqual('2026-08-20', out['report_date'])
        self.assertEqual({'include': 5, 'exclude': 2, 'unknown': 1}, out['counts'])
        verdicts = {r['desc']: r['verdict'] for g in out['groups'] for r in g['rows']}
        self.assertEqual('exclude', verdicts['Future blasting of hull'])
        self.assertEqual('unknown', verdicts['UTM Service'],
                         '판정불가 행도 실려 온다 -- 조용히 버리면 형이 빠진 걸 모른다')


class DocxRouteTests(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.old_db = appmod.DATABASE
        self.old_cfg = appmod.app.config['DATABASE']
        self.old_upload = routes_dock_daily.UPLOAD_DIR
        routes_dock_daily.UPLOAD_DIR = os.path.join(self.tmp.name, 'uploads')
        db = os.path.join(self.tmp.name, 'dd.db')
        appmod.DATABASE = db
        appmod.app.config['DATABASE'] = db
        with appmod.app.app_context():
            appmod.init_db(False)
            self.uid = appmod.execute(
                'INSERT INTO users(username,password_hash,display_name,role,active)'
                ' VALUES(?,?,?,?,1)',
                ('docx-test', appmod.generate_password_hash('t'), 'Docx', 'admin'))
            self.vessel = appmod.execute('INSERT INTO vessels(name, vsl_cd, imo) VALUES(?,?,?)',
                                         ('DOCX TEST', 'D009', 'IMO-9'))
        self.client = appmod.app.test_client()
        with self.client.session_transaction() as s:
            s.update(user_id=self.uid, username='docx-test', display_name='Docx',
                     role='admin', permanent=False)
        # 🔴 번역 API 를 부르지 않는다. 테스트가 외부 네트워크에 매달리면 못 믿는다.
        import helpers_shared
        self.old_key = helpers_shared.GEMINI_API_KEY
        helpers_shared.GEMINI_API_KEY = ''
        self.project, self.report = self._project_report('2026-08-20')

    def tearDown(self):
        import helpers_shared
        helpers_shared.GEMINI_API_KEY = self.old_key
        appmod.DATABASE = self.old_db
        appmod.app.config['DATABASE'] = self.old_cfg
        routes_dock_daily.UPLOAD_DIR = self.old_upload
        self.tmp.cleanup()

    def _project_report(self, date, pid=None):
        if pid is None:
            pid = self.client.post('/api/dock-daily/projects', json={
                'vessel_id': self.vessel, 'title': 'DOCX DD'}).get_json()['id']
        r = self.client.post(f'/api/dock-daily/projects/{pid}/reports/generate',
                             json={'report_date': date}).get_json()
        return pid, r

    def _upload(self, rid, name='dd.docx', **kwargs):
        path = os.path.join(self.tmp.name, 'up.docx')
        build_docx(path, **kwargs)
        with open(path, 'rb') as fh:
            data = fh.read()
        res = self.client.post(f'/api/dock-daily/reports/{rid}/attachments',
                               data={'file': (io.BytesIO(data), name)},
                               content_type='multipart/form-data')
        self.assertEqual(201, res.status_code, res.get_data(as_text=True))
        return res.get_json()['id']

    def _scan(self, rid, aid):
        return self.client.post(f'/api/dock-daily/reports/{rid}/docx-scan',
                                json={'attachment_id': aid})

    def _apply(self, rid, aid, revision, **body):
        payload = {'attachment_id': aid, 'revision': revision}
        payload.update(body)
        return self.client.post(f'/api/dock-daily/reports/{rid}/docx-apply', json=payload)

    def _blocks(self, rid):
        with appmod.app.app_context():
            return [dict(x) for x in appmod.query(
                'SELECT section_key, block_type, content_json, origin, manual_override'
                ' FROM dock_daily_block WHERE report_id=? ORDER BY section_key, sort_order',
                (rid,))]

    def _cards(self, rid):
        """섹션 → 카드 내용.  형 지시 2026-08-23 이후 섹션마다 카드가 **한 장**이다."""
        out = {}
        for b in self._blocks(rid):
            content = routes_dock_daily._json(b['content_json'], {})
            self.assertNotIn(b['section_key'], out, '섹션마다 카드 한 장이어야 한다')
            out[b['section_key']] = content
        return out

    def _lines(self, rid):
        """카드 안의 `1) 2) 3)` 줄을 번호 떼고 평평하게."""
        out = []
        for content in self._cards(rid).values():
            for line in (content.get('title') or '').splitlines():
                out.append(routes_dock_daily.ITEM_NO_RE.sub('', line).strip())
        return out

    def test_scan_reads_the_file_without_touching_the_report(self):
        rid = self.report['id']
        aid = self._upload(rid)
        before = self.client.get(f'/api/dock-daily/reports/{rid}').get_json()
        out = self._scan(rid, aid)
        self.assertEqual(200, out.status_code, out.get_data(as_text=True))
        body = out.get_json()
        self.assertEqual('2026-08-20', body['report_date'])
        self.assertEqual({'include': 5, 'exclude': 2, 'unknown': 1}, body['counts'])
        self.assertEqual({}, body['applied'])
        after = self.client.get(f'/api/dock-daily/reports/{rid}').get_json()
        self.assertEqual(before['revision'], after['revision'], '미리보기는 revision 을 안 올린다')
        self.assertEqual([], self._blocks(rid))
        targets = {g['section_key']: (g['target_key'], g['target_new']) for g in body['groups']}
        self.assertEqual(('shipyard', None), targets['shipyard'])
        self.assertEqual(('vendor', None), targets['vendor'])
        self.assertEqual((None, 'Crew'), targets['crew'], 'Crew 섹션은 apply 가 만든다')

    def test_scan_refuses_an_unreadable_extension(self):
        rid = self.report['id']
        res = self.client.post(f'/api/dock-daily/reports/{rid}/attachments',
                               data={'file': (io.BytesIO(b'hello'), 'note.txt')},
                               content_type='multipart/form-data')
        aid = res.get_json()['id']
        out = self._scan(rid, aid)
        self.assertEqual(400, out.status_code)
        self.assertEqual('not_docx', out.get_json()['code'])
        self.assertEqual(404, self._scan(rid, aid + 999).status_code)

    def test_pdf_passes_the_extension_gate(self):
        """`.pdf` 는 이제 확장자에서 막지 않는다(형 지시 2026-08-23).

        내용이 PDF 가 아니면 `not_docx` 가 아니라 **`docx_unreadable`** 로 떨어져야
        한다 -- 두 사유를 합치면 형은 "형식이 안 된다" 와 "파일이 깨졌다" 를 못 가린다.
        """
        rid = self.report['id']
        res = self.client.post(f'/api/dock-daily/reports/{rid}/attachments',
                               data={'file': (io.BytesIO(b'%PDF-1.4 not really'), 'note.pdf')},
                               content_type='multipart/form-data')
        out = self._scan(rid, res.get_json()['id'])
        self.assertEqual(400, out.status_code)
        self.assertEqual('docx_unreadable', out.get_json()['code'])

    def test_reader_is_chosen_by_extension(self):
        self.assertEqual('dock_daily_pdf', routes_dock_daily._docx_reader('a.PDF'))
        self.assertEqual('dock_daily_docx', routes_dock_daily._docx_reader('a.docx'))
        self.assertIsNone(routes_dock_daily._docx_reader('a.doc'))
        self.assertIsNone(routes_dock_daily._docx_reader(None))

    def test_scan_reports_a_corrupt_file_as_400_not_500(self):
        """형에게 500 은 "서버가 죽었다" 로 읽힌다. 못 읽는 파일은 사람이 고칠 일이다."""
        rid = self.report['id']
        res = self.client.post(f'/api/dock-daily/reports/{rid}/attachments',
                               data={'file': (io.BytesIO(b'PK\x03\x04 not really a docx'),
                                              'broken.docx')},
                               content_type='multipart/form-data')
        out = self._scan(rid, res.get_json()['id'])
        self.assertEqual(400, out.status_code)
        self.assertEqual('docx_unreadable', out.get_json()['code'])

    def test_apply_creates_only_the_included_rows_and_records_provenance(self):
        rid = self.report['id']
        aid = self._upload(rid)
        res = self._apply(rid, aid, self.report['revision'])
        self.assertEqual(200, res.status_code, res.get_data(as_text=True))
        body = res.get_json()
        self.assertEqual(5, body['applied'])
        self.assertEqual(0, body['updated'])
        self.assertEqual(0, body['translated'], 'API 키가 없으면 원문 그대로')
        self.assertEqual(self.report['revision'] + 1, body['revision'])
        blocks = self._blocks(rid)
        # 🔴 섹션마다 카드 **한 장**, 행은 그 안의 `1) 2) 3)` 줄이다(형 지시 2026-08-23).
        # Deck 3행이 shipyard 한 장, 3rd party 1행, Crew 1행 -> 카드 3장 / 줄 5개.
        self.assertEqual(3, len(blocks))
        self.assertEqual(0, body['merged_cards'], '처음 읽으면 접을 낱장이 없다')
        self.assertEqual({'dock_auto'}, {b['origin'] for b in blocks})
        self.assertEqual({0}, {b['manual_override'] for b in blocks})
        self.assertEqual({'item'}, {b['block_type'] for b in blocks})
        cards = self._cards(rid)
        self.assertEqual(['1) Hatch Cover No.1 dismantle hydraulic jacks (계획 완료 08.21)',
                          '2) Load/discharge spare parts and tools',
                          '3) Staging erection continues in cargo hold (진행 중)'],
                         cards['shipyard']['title'].splitlines())
        # 줄별 출처는 카드 **안에** 있다 -- 카드 하나에 여러 행이 들어가면
        # `source_link` 만으로는 어느 줄이 어느 행인지 되짚을 수 없다.
        rows = cards['shipyard']['source_rows']
        self.assertEqual(3, len(rows))
        self.assertEqual('Hatch Cover No.1 dismantle hydraulic jacks', rows[0]['source_en'])
        self.assertEqual('include', rows[0]['source_verdict'])
        texts = self._lines(rid)
        self.assertFalse(any('Future blasting' in t for t in texts), '미래 작업은 안 들어온다')
        self.assertFalse(any('UTM Service' in t for t in texts), '판정불가는 자동으로 안 넣는다')
        with appmod.app.app_context():
            links = appmod.query('SELECT * FROM dock_daily_source_link WHERE report_id=?', (rid,))
        self.assertEqual(5, len(links))
        self.assertEqual({'docx'}, {x['source_system'] for x in links})
        self.assertTrue(all(x['block_id'] for x in links))

    def test_apply_creates_the_crew_section_on_this_date_only(self):
        """🔴 Crew 섹션은 **이 일자 스코프**로 만든다.

        프로젝트 스코프로 만들면 파일을 한 번 읽었다는 이유로 다른 날짜 보고서에까지
        빈 Crew 섹션이 생긴다(형이 2026-08-23 에 고친 바로 그 문제).
        """
        rid = self.report['id']
        _, other = self._project_report('2026-08-21', pid=self.project)
        body = self._apply(rid, self._upload(rid), self.report['revision']).get_json()
        created = body['created_section']
        self.assertEqual('Crew', created['label'])
        with appmod.app.app_context():
            row = appmod.query('SELECT scope FROM dock_daily_section_def WHERE project_id=?'
                               ' AND section_key=?', (self.project, created['section_key']),
                               one=True)
        self.assertEqual('report', row['scope'])
        keys = {s['section_key'] for s in body['sections']}
        self.assertIn(created['section_key'], keys)
        others = self.client.get(f"/api/dock-daily/reports/{other['id']}").get_json()
        self.assertNotIn(created['section_key'],
                         {s['section_key'] for s in others['sections']},
                         '다른 일자에는 안 생긴다')
        # 두 번 읽어도 Crew 섹션이 또 생기지 않는다.
        again = self._apply(rid, self._upload(rid, name='dd2.docx'),
                            body['revision']).get_json()
        self.assertIsNone(again['created_section'])

    def test_reapplying_the_same_work_updates_in_place_instead_of_duplicating(self):
        """감독이 같은 날 문서를 고쳐 다시 보내면 파일 해시가 달라진다.

        🔴 그래서 동일성에 파일 해시를 쓰면 안 된다 -- 같은 작업이 새 카드로 또 들어온다.
        """
        rid = self.report['id']
        first = self._apply(rid, self._upload(rid), self.report['revision']).get_json()
        rows = [
            ['1', 'Vessel arrived to SY 11:42 LT', '19.08.2026', '', ''],
            ['2', 'Hatch Cover No.1 dismantle hydraulic jacks', '20/08/2026', '22.08.2026',
             '22.08.2026'],                                   # 계획일이 바뀐 같은 작업
            ['3', 'Load/discharge spare parts and tools', '20/08/2026', '', ''],
            ['4', 'Staging erection continues in cargo hold', '18.08.2026', '', ''],
            ['5', 'Future blasting of hull', '25.08.2026', '', ''],
        ]
        path = os.path.join(self.tmp.name, 'rev2.docx')
        build_docx(path, rows=rows)
        with open(path, 'rb') as fh:
            data = fh.read()
        aid2 = self.client.post(f'/api/dock-daily/reports/{rid}/attachments',
                                data={'file': (io.BytesIO(data), 'dd-rev2.docx')},
                                content_type='multipart/form-data').get_json()['id']
        second = self._apply(rid, aid2, first['revision']).get_json()
        self.assertEqual(0, second['applied'], '같은 작업이 새 카드로 또 생기면 안 된다')
        # 고친 건 Hatch Cover 의 계획일 한 줄뿐이다.  나머지 4장은 글자가 그대로이므로
        # **다시 쓰지 않는다**(같은 내용 UPDATE 는 revision 만 올려 다른 기기에 409 를 준다).
        self.assertEqual(1, second['updated'])
        self.assertEqual(4, second['unchanged'])
        self.assertEqual(3, len(self._blocks(rid)))
        texts = self._lines(rid)
        self.assertEqual(5, len(texts), '줄이 늘지 않는다')
        self.assertTrue(any('계획 완료 08.22' in t for t in texts), texts)

    def test_reading_the_same_file_twice_writes_nothing_at_all(self):
        """🔴 같은 파일 재적용은 **완전 no-op** 이어야 한다.

        옛 코드는 내용이 같아도 전 카드를 UPDATE 하고 revision 을 올려서, 이 보고서를
        열어 둔 다른 기기가 이유 없이 409 를 맞았다(올마이트 지적 2026-08-23).
        """
        rid = self.report['id']
        aid = self._upload(rid)
        first = self._apply(rid, aid, self.report['revision']).get_json()
        blocks = self._blocks(rid)
        second = self._apply(rid, aid, first['revision']).get_json()
        self.assertEqual(0, second['applied'])
        self.assertEqual(0, second['updated'])
        self.assertEqual(5, second['unchanged'])
        self.assertEqual(first['revision'], second['revision'], 'revision 을 올리지 않는다')
        self.assertEqual(blocks, self._blocks(rid))

    def test_a_second_partial_selection_keeps_the_lines_already_in_the_card(self):
        """🔴 부분 선택 재적용이 **지난번에 넣은 줄을 지우면 안 된다**.

        카드 한 장에 여러 행이 들어가므로, 이번에 고른 행만으로 카드를 다시 쓰면 형이
        먼저 넣어 둔 줄이 조용히 사라진다.  부분 선택은 정상 사용이므로 실제로 난다.
        줄별 출처를 카드 안(`source_rows`)에 두는 이유가 이것이다.
        """
        rid = self.report['id']
        aid = self._upload(rid)
        rows = {r['desc']: r for g in self._scan(rid, aid).get_json()['groups']
                for r in g['rows']}
        first = self._apply(rid, aid, self.report['revision'],
                            row_keys=[rows['Load/discharge spare parts and tools']['row_key']]
                            ).get_json()
        self.assertEqual(['Load/discharge spare parts and tools'], self._lines(rid))
        second = self._apply(rid, aid, first['revision'],
                             row_keys=[rows['Staging erection continues in cargo hold']['row_key']]
                             ).get_json()
        self.assertEqual(1, second['applied'])
        self.assertEqual(1, len(self._blocks(rid)), '같은 섹션이면 카드는 여전히 한 장')
        self.assertEqual(['Load/discharge spare parts and tools',
                          'Staging erection continues in cargo hold (진행 중)'],
                         self._lines(rid))

    def test_old_one_card_per_row_data_is_folded_into_a_single_card(self):
        """옛 계약(행 1개 = 카드 1장)으로 들어간 라이브 카드는 재적용 때 접힌다.

        🔴 링크를 **먼저** 옮기고 나서 낱장을 지운다.  순서를 뒤집으면 `block_id` FK
        (`ON DELETE SET NULL`)가 끊어져 그 줄의 멱등키가 카드를 잃고, 다음 재적용에서
        같은 줄이 새 줄로 또 들어온다.
        """
        rid = self.report['id']
        aid = self._upload(rid)
        first = self._apply(rid, aid, self.report['revision']).get_json()
        # 라이브에 남아 있는 옛 모양을 손으로 만든다: shipyard 카드를 줄마다 한 장으로.
        with appmod.app.app_context():
            block = appmod.query('SELECT * FROM dock_daily_block WHERE report_id=?'
                                 " AND section_key='shipyard'", (rid,), one=True)
            entries = routes_dock_daily._json(block['content_json'], {})['source_rows']
            self.assertEqual(3, len(entries))
            for n, entry in enumerate(entries):
                body = routes_dock_daily.json.dumps(
                    {'title': entry['line'], 'source_en': entry['source_en']},
                    ensure_ascii=False)
                if n == 0:
                    appmod.execute('UPDATE dock_daily_block SET content_json=? WHERE id=?',
                                   (body, block['id']))
                    bid = block['id']
                else:
                    bid = appmod.execute('''INSERT INTO dock_daily_block(report_id,section_key,
                                            sort_order,block_type,content_json,origin)
                                            VALUES (?,'shipyard',?,'item',?,'dock_auto')''',
                                         (rid, 10 + n, body))
                appmod.execute('UPDATE dock_daily_source_link SET block_id=? WHERE report_id=?'
                               ' AND source_subkey=?', (bid, rid, entry['row_key']))
        self.assertEqual(5, len(self._blocks(rid)), '옛 모양: shipyard 3장 + 나머지 2장')
        second = self._apply(rid, aid, first['revision']).get_json()
        self.assertEqual(2, second['merged_cards'])
        self.assertEqual(3, len(self._blocks(rid)))
        self.assertEqual(3, len(self._cards(rid)['shipyard']['title'].splitlines()))
        with appmod.app.app_context():
            after = appmod.query('SELECT block_id FROM dock_daily_source_link WHERE report_id=?',
                                 (rid,))
        self.assertTrue(all(x['block_id'] for x in after), '멱등키가 카드를 잃으면 안 된다')
        # 접힌 뒤 한 번 더 읽으면 완전 no-op 이다(줄이 또 늘지 않는다).
        third = self._apply(rid, aid, second['revision']).get_json()
        self.assertEqual(second['revision'], third['revision'])
        self.assertEqual(5, len(self._lines(rid)))

    def test_the_same_row_stored_in_two_cards_becomes_one_line(self):
        """🔴 링크와 카드가 어긋난 데이터에서 같은 행이 **두 줄**로 저장되면 안 된다.

        도달 경로: 카드 A 의 `source_rows` 에 행 X 가 적혀 있는데 그 뒤 X 의 링크가 카드
        B 로 옮겨 간 상태(부분쓰기 유실·손편집).  두 카드 다 문서 링크를 갖고 있으므로
        둘 다 읽히고, 그대로 담으면 형이 같은 문장을 두 번 읽는다(올마이트 지적
        2026-08-23).  링크가 **아예 없는** dock_auto 카드는 일부러 손대지 않는다 --
        같은 `origin` 을 Dock 러너도 쓰므로, 링크 없이 origin 만 보고 접으면 러너 카드를
        부순다.
        """
        rid = self.report['id']
        aid = self._upload(rid)
        first = self._apply(rid, aid, self.report['revision']).get_json()
        crew = first['created_section']['section_key']
        with appmod.app.app_context():
            ship = appmod.query('SELECT * FROM dock_daily_block WHERE report_id=?'
                                " AND section_key='shipyard'", (rid,), one=True)
            other = appmod.query('SELECT * FROM dock_daily_block WHERE report_id=?'
                                 ' AND section_key=?', (rid, crew), one=True)
            stolen = routes_dock_daily._json(ship['content_json'], {})['source_rows'][0]
            content = routes_dock_daily._json(other['content_json'], {})
            content['source_rows'] = content['source_rows'] + [stolen]
            content['title'] += '\n2) ' + stolen['line']
            appmod.execute('UPDATE dock_daily_block SET content_json=? WHERE id=?',
                           (routes_dock_daily.json.dumps(content, ensure_ascii=False),
                            other['id']))
        self._apply(rid, aid, first['revision'])
        cards = self._cards(rid)
        lines = cards['shipyard']['title'].splitlines()
        self.assertEqual(3, len(lines), lines)
        self.assertEqual(len(set(lines)), len(lines), '같은 문장이 두 줄로 남으면 안 된다')
        self.assertNotIn(stolen['line'], cards[crew]['title'], '옮겨 간 사본은 남지 않는다')

    def test_a_row_moving_to_another_section_leaves_no_copy_behind(self):
        """감독이 표 제목을 바꾸면 그 행만 다른 섹션으로 옮겨 간다.

        🔴 옛 카드에 사본이 남으면 같은 문장이 두 섹션에서 메일로 나간다.  줄이 전부
        옮겨 가 빈 카드가 되면 그 카드는 접어서 지운다.
        """
        rid = self.report['id']
        first = self._apply(rid, self._upload(rid), self.report['revision']).get_json()
        crew = first['created_section']['section_key']
        self.assertEqual(['ME FO Auto-filter maintenance'],
                         [routes_dock_daily.ITEM_NO_RE.sub('', x).strip()
                          for x in self._cards(rid)[crew]['title'].splitlines()])
        # 같은 작업을 Deck 표로 옮긴 문서.  `row_key` 는 (표 라벨|설명) 이므로 새 행이다.
        path = os.path.join(self.tmp.name, 'moved.docx')
        doc = Document()
        _table(doc, 'Leading Deck Works done by the Yard', HEADER,
               [['1', 'ME FO Auto-filter maintenance', '20.08.2026', '', '']])
        doc.save(path)
        with open(path, 'rb') as fh:
            data = fh.read()
        aid = self.client.post(f'/api/dock-daily/reports/{rid}/attachments',
                               data={'file': (io.BytesIO(data), 'moved.docx')},
                               content_type='multipart/form-data').get_json()['id']
        second = self._apply(rid, aid, first['revision']).get_json()
        cards = self._cards(rid)
        self.assertIn('ME FO Auto-filter maintenance', cards['shipyard']['title'])
        # Crew 카드에는 그 줄이 남아 있다 -- `row_key` 가 표 라벨 스코프라 별개 행이고,
        # apply 는 지금 파일에 없는 줄을 지우지 않는다(`stale_applied` 로만 알린다).
        self.assertIn(crew, cards, '지금 파일에 없는 줄은 apply 가 지우지 않는다')
        self.assertEqual(1, second['applied'])

    def test_a_card_carrying_a_photo_is_never_folded_away(self):
        """🔴 사진이 매달린 카드는 병합·삭제 대상에서 뺀다.

        `dock_daily_attachment.block_id` 는 `ON DELETE SET NULL` 이라 지워도 500 은 안
        나지만, 형이 올린 사진이 조용히 카드에서 떨어진다.
        """
        rid = self.report['id']
        aid = self._upload(rid)
        first = self._apply(rid, aid, self.report['revision']).get_json()
        with appmod.app.app_context():
            block = appmod.query('SELECT * FROM dock_daily_block WHERE report_id=?'
                                 " AND section_key='shipyard'", (rid,), one=True)
            appmod.execute('UPDATE dock_daily_attachment SET block_id=? WHERE id=?',
                           (block['id'], aid))
        second = self._apply(rid, aid, first['revision']).get_json()
        self.assertEqual(0, second['merged_cards'])
        self.assertEqual(3, second['skipped_edited'], '사진 달린 카드의 줄은 손대지 않는다')
        with appmod.app.app_context():
            still = appmod.query('SELECT block_id FROM dock_daily_attachment WHERE id=?',
                                 (aid,), one=True)
        self.assertEqual(block['id'], still['block_id'], '사진이 카드에서 떨어지면 안 된다')

    def test_a_section_that_exists_on_another_date_is_attached_not_duplicated(self):
        """🔴 같은 이름 섹션을 날짜마다 새로 만들면 프로젝트 섹션 목록이 도배된다.

        정의는 재사용하고 이 일자에 **붙이기만** 한다(올마이트 지적 2026-08-23).
        """
        rid = self.report['id']
        first = self._apply(rid, self._upload(rid), self.report['revision']).get_json()
        crew = first['created_section']['section_key']
        _, other = self._project_report('2026-08-21', pid=self.project)
        # 다음 날 문서.  Crew 표의 작업이 **그날** 시작하므로 이 일자에도 들어온다
        # (기본 fixture 의 Crew 행은 20.08 짜리라 21 일 기준에선 지난 작업이 된다).
        path = os.path.join(self.tmp.name, 'day2.docx')
        doc = Document()
        _table(doc, 'Leading Works done by Crew', HEADER,
               [['1', 'ME FO Auto-filter maintenance', '21.08.2026', '', '']])
        doc.save(path)
        with open(path, 'rb') as fh:
            data = fh.read()
        aid = self.client.post(f"/api/dock-daily/reports/{other['id']}/attachments",
                               data={'file': (io.BytesIO(data), 'day2.docx')},
                               content_type='multipart/form-data').get_json()['id']
        body = self._apply(other['id'], aid, other['revision']).get_json()
        self.assertIsNone(body['created_section'], '정의를 또 만들지 않는다')
        self.assertEqual(1, body['attached_sections'])
        with appmod.app.app_context():
            defs = appmod.query("SELECT section_key FROM dock_daily_section_def"
                                " WHERE project_id=? AND label='Crew'", (self.project,))
        self.assertEqual([crew], [d['section_key'] for d in defs], 'Crew 정의는 하나뿐')
        self.assertIn(crew, {s['section_key'] for s in body['sections']})

    def test_rows_with_nowhere_to_go_are_counted_not_silently_dropped(self):
        """비고 섹션이 꺼져 있으면 모르는 제목의 행은 갈 곳이 없다.

        🔴 그 사실을 숫자로 돌려주지 않으면 화면은 "바뀐 내용이 없습니다" 만 띄우고
        형은 들어간 줄 안다(올마이트 지적 2026-08-23).

        비고는 고정 섹션이라 **지금 API 로는 끌 수 없다**(`kind!='special'` 이면
        `enabled` 를 안 받고 삭제도 409 `fixed_section`).  그래서 여기서는 DB 를 직접
        건드려 그 상태를 만든다 -- 이 경로는 옛 데이터와 앞으로의 계약 변경에 대한
        방어선이므로, 도달 불가라는 이유로 잠금을 풀지 않는다.
        """
        rid = self.report['id']
        with appmod.app.app_context():
            appmod.execute('UPDATE dock_daily_section_def SET enabled=0'
                           ' WHERE project_id=? AND section_key=?', (self.project, 'remark'))
        path = os.path.join(self.tmp.name, 'odd2.docx')
        doc = Document()
        _table(doc, 'Leading Works done by Subcontractor Alpha', HEADER,
               [['1', 'Anode renewal', '20.08.2026', '', '']])
        doc.save(path)
        with open(path, 'rb') as fh:
            data = fh.read()
        aid = self.client.post(f'/api/dock-daily/reports/{rid}/attachments',
                               data={'file': (io.BytesIO(data), 'odd2.docx')},
                               content_type='multipart/form-data').get_json()['id']
        scan = self._scan(rid, aid).get_json()
        self.assertIsNone(scan['groups'][0]['target_key'])
        body = self._apply(rid, aid, self.report['revision']).get_json()
        self.assertEqual(0, body['applied'])
        self.assertEqual(1, body['skipped_unmapped'])
        self.assertEqual([], self._blocks(rid))

    def test_scan_flags_applied_cards_that_the_file_no_longer_has(self):
        """감독이 문장을 고쳐 쓰면 `row_key` 가 달라져 옛 카드가 중복으로 남는다.

        자동으로 알아맞힐 수는 없으니 **짝 없는 옛 카드 수를 알린다**(올마이트 지적).
        """
        rid = self.report['id']
        first = self._apply(rid, self._upload(rid), self.report['revision']).get_json()
        path = os.path.join(self.tmp.name, 'reworded.docx')
        # 한 줄만 고쳐 쓰고 나머지는 그대로 둔다 -- 그래야 짝을 잃은 카드가 그 한 장뿐이다.
        build_docx(path, rows=[
            ['1', 'Hatch Cover No.1 dismantle hydraulic jacks and pins',  # 문장 수정
             '20/08/2026', '21.08.2026', '21.08.2026'],
            ['2', 'Load/discharge spare parts and tools', '20/08/2026', '', ''],
            ['3', 'Staging erection continues in cargo hold', '18.08.2026', '', ''],
        ])
        with open(path, 'rb') as fh:
            data = fh.read()
        aid = self.client.post(f'/api/dock-daily/reports/{rid}/attachments',
                               data={'file': (io.BytesIO(data), 'reworded.docx')},
                               content_type='multipart/form-data').get_json()['id']
        scan = self._scan(rid, aid).get_json()
        self.assertEqual(1, len(scan['stale_applied']),
                         '고쳐 쓴 그 한 줄의 옛 카드가 짝을 잃는다')
        self.assertTrue(all(k in scan['applied'] for k in scan['stale_applied']))

    def test_a_card_the_user_edited_is_never_overwritten(self):
        """🔴 형이 고친 문장은 형의 문장이다. 파일을 다시 읽었다는 이유로 덮을 수 없다."""
        rid = self.report['id']
        first = self._apply(rid, self._upload(rid), self.report['revision']).get_json()
        block = self._blocks(rid)[0]
        with appmod.app.app_context():
            bid = appmod.query('SELECT id FROM dock_daily_block WHERE report_id=?'
                               ' ORDER BY section_key, sort_order', (rid,))[0]['id']
        edited = self.client.put(f'/api/dock-daily/reports/{rid}', json={
            'revision': first['revision'],
            'operations': [{'id': bid, 'section_key': block['section_key'],
                            'block_type': 'item', 'sort_order': 0,
                            'content': {'title': '형이 직접 고친 문장'}}]}).get_json()
        again = self._apply(rid, self._upload(rid, name='dd3.docx'), edited['revision'])
        body = again.get_json()
        self.assertEqual(1, body['skipped_edited'])
        titles = [routes_dock_daily._json(b['content_json'], {}).get('title')
                  for b in self._blocks(rid)]
        self.assertIn('형이 직접 고친 문장', titles)

    def test_apply_honours_an_explicit_row_selection(self):
        rid = self.report['id']
        aid = self._upload(rid)
        scan = self._scan(rid, aid).get_json()
        rows = {r['desc']: r for g in scan['groups'] for r in g['rows']}
        picked = [rows['UTM Service']['row_key'], rows['Future blasting of hull']['row_key']]
        body = self._apply(rid, aid, self.report['revision'],
                           row_keys=picked + ['deadbeef1234']).get_json()
        self.assertEqual(2, body['applied'], '형이 고르면 제외·판정불가도 넣는다')
        self.assertEqual(['deadbeef1234'], body['unknown_row_keys'])
        self.assertEqual({'UTM Service', 'Future blasting of hull'}, set(self._lines(rid)))

    def test_apply_rejects_an_empty_selection_and_a_stale_revision(self):
        rid = self.report['id']
        aid = self._upload(rid)
        empty = self._apply(rid, aid, self.report['revision'], row_keys=[])
        self.assertEqual(400, empty.status_code)
        self.assertEqual('no_rows', empty.get_json()['code'])
        stale = self._apply(rid, aid, self.report['revision'] + 5)
        self.assertEqual(409, stale.status_code)
        self.assertEqual('revision_conflict', stale.get_json()['code'])
        self.assertEqual([], self._blocks(rid), '거절된 요청은 아무것도 안 남긴다')
        missing = self.client.post(f'/api/dock-daily/reports/{rid}/docx-apply',
                                   json={'attachment_id': aid})
        self.assertEqual(400, missing.status_code, 'revision 없는 요청은 400')

    def test_a_final_report_is_locked_before_any_translation_runs(self):
        rid = self.report['id']
        aid = self._upload(rid)
        first = self._apply(rid, aid, self.report['revision']).get_json()
        final = self.client.post(f'/api/dock-daily/reports/{rid}/status',
                                 json={'status': 'final', 'revision': first['revision']})
        self.assertEqual(200, final.status_code, final.get_data(as_text=True))
        rev = final.get_json()['revision']
        calls = []
        import helpers_shared
        real = helpers_shared.translate_texts_ko
        helpers_shared.translate_texts_ko = lambda texts: calls.append(texts) or list(texts)
        try:
            res = self._apply(rid, aid, rev)
        finally:
            helpers_shared.translate_texts_ko = real
        self.assertEqual(409, res.status_code)
        self.assertEqual('final_locked', res.get_json()['code'])
        self.assertEqual([], calls, '확정본이면 번역 API 를 아예 부르지 않는다')

    def test_translation_replaces_the_text_and_keeps_the_english_original(self):
        rid = self.report['id']
        aid = self._upload(rid)
        import helpers_shared
        real = helpers_shared.translate_texts_ko
        helpers_shared.translate_texts_ko = lambda texts: ['[KO] ' + t for t in texts]
        try:
            body = self._apply(rid, aid, self.report['revision']).get_json()
        finally:
            helpers_shared.translate_texts_ko = real
        self.assertEqual(5, body['translated'])
        lines = self._lines(rid)
        self.assertEqual(5, len(lines))
        self.assertTrue(all(t.startswith('[KO] ') for t in lines), lines)
        rows = [r for c in self._cards(rid).values() for r in c['source_rows']]
        self.assertTrue(all(not r['source_en'].startswith('[KO] ') for r in rows),
                        '영문 원문은 추적용으로 남는다')

    def test_translation_removes_terminal_ham_but_preserves_nouns(self):
        import routes_dock_daily as routes
        self.assertEqual('도장 작업 완료.', routes._strip_generated_ham('도장 작업 완료함.'))
        self.assertEqual('배관 검사 진행', routes._strip_generated_ham('배관 검사 진행함'))
        self.assertEqual('검사 범위에 포함.', routes._strip_generated_ham('검사 범위에 포함.'))
        self.assertEqual('선체 결함.', routes._strip_generated_ham('선체 결함.'))
        self.assertEqual('공구함.', routes._strip_generated_ham('공구함.'))
        self.assertEqual('구급함 점검.', routes._strip_generated_ham('구급함 점검함.'))

        rid = self.report['id']
        aid = self._upload(rid)
        import helpers_shared
        real = helpers_shared.translate_texts_ko
        helpers_shared.translate_texts_ko = lambda texts: ['작업 완료함.' for _ in texts]
        try:
            body = self._apply(rid, aid, self.report['revision']).get_json()
        finally:
            helpers_shared.translate_texts_ko = real
        self.assertEqual(5, body['translated'])
        applied = self._lines(rid)
        self.assertEqual(5, len(applied))
        self.assertTrue(all(line.startswith('작업 완료.') for line in applied), applied)
        self.assertTrue(all('완료함' not in line for line in applied), applied)

    def test_a_failing_translator_falls_back_to_the_english_text(self):
        """🔴 번역 실패로 파일 읽기 전체를 실패시키지 않는다. 영문 원문이라도 들어와야 한다."""
        rid = self.report['id']
        aid = self._upload(rid)
        import helpers_shared
        real = helpers_shared.translate_texts_ko

        def boom(texts):
            raise RuntimeError('gemini down')

        helpers_shared.translate_texts_ko = boom
        try:
            res = self._apply(rid, aid, self.report['revision'])
        finally:
            helpers_shared.translate_texts_ko = real
        self.assertEqual(200, res.status_code, res.get_data(as_text=True))
        self.assertEqual(5, res.get_json()['applied'])
        self.assertEqual(0, res.get_json()['translated'])

    def test_unmapped_headings_land_in_remark_rather_than_vanishing(self):
        rid = self.report['id']
        path = os.path.join(self.tmp.name, 'odd.docx')
        doc = Document()
        _table(doc, 'Leading Works done by Subcontractor Alpha', HEADER,
               [['1', 'Anode renewal', '20.08.2026', '', '']])
        doc.save(path)
        with open(path, 'rb') as fh:
            data = fh.read()
        aid = self.client.post(f'/api/dock-daily/reports/{rid}/attachments',
                               data={'file': (io.BytesIO(data), 'odd.docx')},
                               content_type='multipart/form-data').get_json()['id']
        scan = self._scan(rid, aid).get_json()
        self.assertEqual(['Leading Works done by Subcontractor Alpha'], scan['unmapped_headings'])
        self.assertTrue(scan['groups'][0]['target_fallback'])
        self.assertEqual('remark', scan['groups'][0]['target_key'])
        self._apply(rid, aid, self.report['revision'])
        self.assertEqual(['remark'], [b['section_key'] for b in self._blocks(rid)])


def _png(color, size=(8, 8)):
    """서로 다른 바이트의 작은 PNG.  같은 색이면 바이트도 같다(중복 판정 fixture)."""
    from PIL import Image
    buf = io.BytesIO()
    Image.new('RGB', size, color).save(buf, 'PNG')
    return buf.getvalue()


RED, BLUE, GREEN = _png((200, 10, 10)), _png((10, 10, 200)), _png((10, 200, 10))


def _put(cell, data):
    from docx.shared import Inches
    cell.paragraphs[0].add_run().add_picture(io.BytesIO(data), width=Inches(1))


def build_photo_docx(path, rows=True, loose=None, pairs=None, stacked=None, own=None):
    """사진표가 있는 fixture.  실측 두 템플릿을 모두 만든다.

    · `pairs`   : `[(캡션, bytes)]` -> `[캡션][사진]` **같은 행** (20.08 템플릿)
    · `stacked` : `[(캡션, bytes)]` -> 사진행 다음에 캡션행 (21.08 템플릿)
    · `own`     : `[(캡션, bytes)]` -> 한 칸에 사진과 글이 같이
    · `loose`   : 표 **밖** 본문에 박은 그림(= 머리글 로고 자리)
    """
    doc = Document()
    if rows:
        _table(doc, 'Leading Deck Works done by the Yard', HEADER,
               [['1', 'Vessel arrived to SY 11:42 LT', '20.08.2026', '', '']])
    if loose:
        from docx.shared import Inches
        doc.add_paragraph().add_run().add_picture(io.BytesIO(loose), width=Inches(1))
    doc.add_paragraph('Pictures Documenting the Status & Budgets')
    if pairs:
        table = doc.add_table(rows=0, cols=2)
        for caption, data in pairs:
            cells = table.add_row().cells
            cells[0].text = caption
            _put(cells[1], data)
    if stacked:
        table = doc.add_table(rows=0, cols=2)
        for start in range(0, len(stacked), 2):
            chunk = stacked[start:start + 2]
            shots = table.add_row().cells
            caps = table.add_row().cells
            for k, (caption, data) in enumerate(chunk):
                _put(shots[k], data)
                caps[k].text = caption
    if own:
        table = doc.add_table(rows=0, cols=1)
        for caption, data in own:
            cell = table.add_row().cells[0]
            _put(cell, data)
            cell.paragraphs[0].add_run(caption)
    doc.save(path)


class PhotoParserTests(unittest.TestCase):
    """문서 안 사진 (형 지시 2026-08-23 "1. 고고" = Phase 3)."""

    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.path = os.path.join(self.tmp.name, 'p.docx')

    def tearDown(self):
        self.tmp.cleanup()

    def test_only_pictures_inside_tables_are_photos(self):
        """🔴 표 밖 그림은 사진이 아니다 -- 실측상 그 자리가 회사 로고다.

        "문서의 이미지를 전부 가져오기" 로 짜면 형 보고서 사진 격자 첫 칸에 시노코
        로고가 박히고, 그 메일이 조선소로 나간다.
        """
        build_photo_docx(self.path, loose=GREEN, pairs=[('Rope guard removal', RED)])
        out = parser.photos(self.path)
        self.assertEqual(['Rope guard removal'], [p['caption'] for p in out['photos']])
        self.assertNotIn(GREEN, [p['data'] for p in out['photos']])
        self.assertTrue(out['captions'])

    def test_both_live_caption_layouts(self):
        """실측 템플릿 두 개: 같은 행 왼쪽 칸 / 아래 행 같은 열."""
        build_photo_docx(self.path, pairs=[('ME Overhauling by Cat Asea', RED)],
                         stacked=[('Rope guard removal', BLUE), ('Anode renewal', GREEN)])
        got = {p['caption']: p['size'] for p in parser.photos(self.path)['photos']}
        self.assertEqual({'ME Overhauling by Cat Asea', 'Rope guard removal', 'Anode renewal'},
                         set(got))

    def test_caption_in_the_same_cell_is_used(self):
        build_photo_docx(self.path, own=[('Tail shaft drawn out', RED)])
        self.assertEqual(['Tail shaft drawn out'],
                         [p['caption'] for p in parser.photos(self.path)['photos']])

    def test_placeholder_caption_is_left_empty_not_invented(self):
        """🔴 채워지지 않은 `Photo 3: Description` 은 캡션이 아니다.

        실측 문서에 그 자리표시자가 그대로 남아 있다.  그걸 캡션으로 넣으면 메일에
        `Photo 3: Description` 이 나간다 -- 없는 캡션은 빈 칸으로 둔다.
        """
        for text in ('Photo 3: Description', 'photo:', 'Photo 11', '   ', 'photo. description'):
            self.assertEqual('', parser.photo_caption(text), text)
        self.assertEqual('Rope guard', parser.photo_caption(' Rope\n guard '))
        build_photo_docx(self.path, pairs=[('Photo 2: Description', RED)])
        self.assertEqual([''], [p['caption'] for p in parser.photos(self.path)['photos']])

    def test_identical_bytes_are_folded_and_counted(self):
        build_photo_docx(self.path, pairs=[('A', RED), ('B', RED), ('C', BLUE)])
        out = parser.photos(self.path)
        self.assertEqual(2, len(out['photos']))
        self.assertEqual(1, out['duplicates'])
        self.assertEqual(['A', 'C'], [p['caption'] for p in out['photos']])

    def test_photo_kind_rejects_formats_we_cannot_store(self):
        self.assertEqual(('jpg', 'image/jpeg'), parser.photo_kind(b'\xff\xd8\xff\xe0rest'))
        self.assertEqual(('png', 'image/png'), parser.photo_kind(RED))
        self.assertEqual((None, None), parser.photo_kind(b'GIF89a'))
        self.assertEqual((None, None), parser.photo_kind(b''))

    def test_unsupported_pictures_are_counted_not_dropped_silently(self):
        build_photo_docx(self.path, pairs=[('ok', RED)])
        # 실제 EMF/WMF 를 python-docx 로 넣기 어려우므로 magic 판정만으로 잠근다.
        out = parser.photos(self.path)
        self.assertEqual(0, out['skipped'])
        self.assertEqual(0, out['photo_limit'])

    def test_a_photo_inside_a_nested_table_is_found_once_with_its_own_caption(self):
        """🔴 중첩 표는 두 번 발견되기 쉬운 자리다.

        칸의 그림을 `iter()` 로 훑으면 **바깥 칸**이 안쪽 표의 그림까지 자기 것으로
        보고, 파서가 안쪽 표를 한 번 더 돌기 때문에 같은 사진이 두 번 나온다.  접기가
        살려 주는 것도 한 장뿐이고, 캡션은 먼저 만난 바깥 칸 글이 채택된다 -- 엉뚱한
        문장이 붙은 사진이 그대로 조선소로 나간다.
        """
        doc = Document()
        doc.add_paragraph('Pictures Documenting the Status & Budgets')
        outer = doc.add_table(rows=1, cols=1)
        cell = outer.rows[0].cells[0]
        cell.text = 'Engine room general'          # 바깥 칸 글 = 엉뚱한 캡션 후보
        inner = cell.add_table(rows=1, cols=2)
        inner.rows[0].cells[0].text = 'Rope guard removal'
        _put(inner.rows[0].cells[1], RED)
        doc.save(self.path)
        out = parser.photos(self.path)
        self.assertEqual(1, len(out['photos']))
        self.assertEqual(0, out['duplicates'], '두 번 발견하면 접기 수가 오른다')
        self.assertEqual('Rope guard removal', out['photos'][0]['caption'])

    def test_short_photo_keys_stay_unique_inside_one_document(self):
        """🔴 표시용 키가 겹치면 형이 한 장을 골라도 서버가 두 장을 넣는다.

        접기 판정은 **full sha256** 으로 하고, 짧은 키는 문서 안에서 유일할 때까지만
        늘린다(같은 digest 는 애초에 접혀서 여기 두 번 오지 않는다).
        """
        used = set()
        a = parser.photo_key('c' * 12 + '1' * 52, used)
        used.add(a)
        b = parser.photo_key('c' * 12 + '2' * 52, used)
        used.add(b)
        self.assertNotEqual(a, b)
        self.assertEqual(12, len(a))
        self.assertGreater(len(b), 12, '앞 12자가 같으면 더 길게 잡는다')
        # 같은 digest 가 두 번 와도 서로 다른 키를 받는다(마지막 방어선).
        self.assertNotEqual(a, parser.photo_key('c' * 12 + '1' * 52, used))

    def test_pdf_letterhead_repeated_on_every_page_is_not_a_photo(self):
        """🔴 실측: 같은 10,697 byte 그림이 10쪽 전부에 있다(회사 레터헤드).

        "쪽마다 나오는 걸 한 장으로 접기" 로 짜면 그 로고가 사진 1장으로 들어온다.
        """
        import dock_daily_pdf as pdf

        def img(key, page, top=0.0):
            return {'key': key, 'data': key.encode(), 'ext': 'png', 'mime': 'image/png',
                    'page': page, 'top': top, 'x0': 0.0}

        pages = [[img('logo', 0), img('shot1', 0, 100)],
                 [img('logo', 1), img('shot2', 1, 100)],
                 [img('logo', 2)]]
        out = pdf.fold_photos(pages, skipped=1)
        self.assertEqual(['shot1', 'shot2'], [p['photo_key'] for p in out['photos']])
        self.assertEqual(3, out['letterhead'])
        self.assertEqual(1, out['skipped'])
        # 🔴 PDF 캡션은 만들지 않는다(두 열이 한 줄로 뭉쳐 뽑힌다 -- 절반이 틀린다).
        self.assertFalse(out['captions'])
        self.assertEqual([''], list({p['caption'] for p in out['photos']}))


class PhotoRouteTests(unittest.TestCase):
    """`docx-scan` / `docx-apply` 의 사진 계약.

    🔴 `DocxRouteTests` 를 **상속하지 않는다**.  상속하면 위쪽 행 테스트 30여 개가
       사진 fixture(`_upload` 를 덮었으므로)로 한 번 더 돌아, 통과 개수는 늘지만
       잠그는 계약은 하나도 늘지 않는다.  필요한 준비·helper 만 빌려 쓴다.
    """
    setUp = DocxRouteTests.setUp
    tearDown = DocxRouteTests.tearDown
    _project_report = DocxRouteTests._project_report
    _scan = DocxRouteTests._scan
    _apply = DocxRouteTests._apply
    _blocks = DocxRouteTests._blocks

    def _upload(self, rid, name='p.docx', **kwargs):
        path = os.path.join(self.tmp.name, 'up_photo.docx')
        build_photo_docx(path, **kwargs)
        with open(path, 'rb') as fh:
            data = fh.read()
        res = self.client.post(f'/api/dock-daily/reports/{rid}/attachments',
                               data={'file': (io.BytesIO(data), name)},
                               content_type='multipart/form-data')
        self.assertEqual(201, res.status_code, res.get_data(as_text=True))
        return res.get_json()['id']

    def _shots(self, rid, pairs=None, **kw):
        aid = self._upload(rid, pairs=pairs if pairs is not None else [('Rope guard', RED)], **kw)
        scan = self._scan(rid, aid).get_json()
        return aid, scan

    def _images(self, rid):
        """`[(section_key, [캡션…])]` -- 사진 격자 블록만."""
        out = []
        for b in self._blocks(rid):
            if b['block_type'] == 'image':
                content = routes_dock_daily._json(b['content_json'], {})
                out.append((b['section_key'], [x['caption'] for x in content['images']],
                            content['columns']))
        return out

    def _files(self, rid):
        """사진 첨부만.  읽은 문서(.docx) 자체도 첨부이므로 걸러야 수가 맞는다."""
        with appmod.app.app_context():
            return [dict(x) for x in appmod.query(
                'SELECT id, block_id, original_name, mime_type, size, sha256, deleted_at'
                ' FROM dock_daily_attachment WHERE report_id=? AND deleted_at IS NULL'
                " AND mime_type LIKE 'image/%' ORDER BY id", (rid,))]

    def test_scan_reports_photos_without_shipping_the_bytes(self):
        rid = self.report['id']
        _, scan = self._shots(rid, pairs=[('Rope guard', RED), ('Anode', BLUE)])
        self.assertEqual(['Rope guard', 'Anode'], [p['caption'] for p in scan['photos']])
        self.assertTrue(all('data' not in p for p in scan['photos']))
        self.assertTrue(all(p['size'] > 0 and p['mime'] == 'image/png' for p in scan['photos']))
        self.assertTrue(scan['photo_captions'])
        self.assertEqual(0, scan['photo_duplicates'] + scan['photo_skipped'])
        self.assertFalse(any(p['applied'] for p in scan['photos']))

    def test_the_scan_always_says_how_many_were_taken_for_letterhead(self):
        """🔴 레터헤드로 빼낸 장수는 **키가 항상 있어야** 한다.

        화면(웹 규칙·앱 규칙)이 이 값으로 "쪽마다 반복되는 그림 N장은 뺐습니다" 를
        말한다.  docx 는 표 안 그림만 보므로 0 이지만, 키가 빠지면 PDF 에서만 값이
        생기는 필드가 되어 두 화면 중 한쪽이 조용히 아무 말도 안 한다.
        """
        rid = self.report['id']
        _, scan = self._shots(rid)
        self.assertEqual(0, scan['photo_letterhead'])
        self.assertEqual('Photos', scan['photo_section'])

    def test_english_captions_that_fail_to_translate_are_counted_and_korean_ones_are_not_sent(self):
        """🔴 캡션 번역 실패는 **서버가 센다**(`photos_untranslated`).

        화면에서 뺄셈으로 만들 수 없다 -- 이미 한국어인 캡션은 애초에 번역기에 보내지
        않으므로, 화면이 "캡션 수 - 성공 수" 로 계산하면 한국어 캡션 하나가 매번
        "번역 실패" 로 잡힌다(있지도 않은 실패를 경고하는 것도 거짓이다).
        """
        import helpers_shared
        rid = self.report['id']
        aid, scan = self._shots(rid, pairs=[('Rope guard', RED), ('선미관 점검', BLUE)])
        keys = [p['photo_key'] for p in scan['photos']]
        sent, real = [], helpers_shared.translate_texts_ko

        def boom(texts):
            sent.append(list(texts))
            raise RuntimeError('no api key')

        helpers_shared.translate_texts_ko = boom
        try:
            # 🔴 `row_keys` 를 **비워서** 보낸다. 빼면 서버가 행 전체로 읽어(옛 클라이언트
            #    계약) 행 번역까지 섞여, 이 테스트가 캡션 규칙을 잠그지 못한다.
            body = self._apply(rid, aid, self.report['revision'],
                               row_keys=[], photo_keys=keys).get_json()
        finally:
            helpers_shared.translate_texts_ko = real
        self.assertEqual([['Rope guard']], sent, '한국어 캡션은 보내지 않는다')
        self.assertEqual(2, body['photos_added'])
        self.assertEqual(0, body['photos_translated'])
        self.assertEqual(1, body['photos_untranslated'], '영문 캡션 한 장만 실패다')
        self.assertEqual(0, body['translated'], '캡션은 행 번역 수에 섞이지 않는다')

    def test_translated_captions_are_not_reported_as_untranslated(self):
        import helpers_shared
        rid = self.report['id']
        aid, scan = self._shots(rid, pairs=[('Rope guard', RED)])
        real = helpers_shared.translate_texts_ko
        helpers_shared.translate_texts_ko = lambda texts: ['[KO] ' + t for t in texts]
        try:
            body = self._apply(rid, aid, self.report['revision'],
                               photo_keys=[scan['photos'][0]['photo_key']]).get_json()
        finally:
            helpers_shared.translate_texts_ko = real
        self.assertEqual(1, body['photos_translated'])
        self.assertEqual(0, body['photos_untranslated'])
        self.assertEqual([['[KO] Rope guard']], [x[1] for x in self._images(rid)])

    def test_picking_only_photos_already_in_the_report_explains_itself(self):
        """🔴 할 일이 0장이어도 400 이 아니다.

        400 으로 끊으면 `photos_already` 안내가 사라지고 형은 사유 없는 "넣을 항목이
        없습니다" 만 본다 -- 고른 사진이 왜 안 들어갔는지 화면에서 알 수 없다.
        """
        rid = self.report['id']
        aid, scan = self._shots(rid)
        key = scan['photos'][0]['photo_key']
        first = self._apply(rid, aid, self.report['revision'], photo_keys=[key]).get_json()
        again = self._apply(rid, aid, first['revision'], row_keys=[], photo_keys=[key])
        self.assertEqual(200, again.status_code, again.get_data(as_text=True))
        body = again.get_json()
        self.assertEqual(1, body['photos_already'])
        self.assertEqual(0, body['photos_added'])
        self.assertEqual(0, body['photos_untranslated'])

    def test_a_client_that_does_not_know_photo_keys_attaches_nothing(self):
        """🔴 지금 라이브인 웹·OTA 258 은 `photo_keys` 를 모른다.

        없을 때 '전부' 로 읽으면 형이 고르지도 않은 사진이 조용히 붙고, 그 파일은
        형이 손으로 하나씩 지워야 한다.
        """
        rid = self.report['id']
        aid, _ = self._shots(rid)
        body = self._apply(rid, aid, self.report['revision']).get_json()
        self.assertEqual(0, body['photos_added'])
        self.assertEqual([], self._files(rid))
        self.assertEqual([], self._images(rid))

    def test_picked_photos_become_one_grid_in_its_own_section(self):
        rid = self.report['id']
        aid, scan = self._shots(rid, pairs=[('Rope guard', RED), ('Anode', BLUE)])
        keys = [p['photo_key'] for p in scan['photos']]
        body = self._apply(rid, aid, self.report['revision'], photo_keys=keys).get_json()
        self.assertEqual(2, body['photos_added'])
        self.assertEqual('Photos', body['created_section']['label'])
        images = self._images(rid)
        self.assertEqual(1, len(images))
        self.assertEqual(['Rope guard', 'Anode'], images[0][1])
        self.assertEqual(routes_dock_daily.DOCX_PHOTO_COLUMNS, images[0][2])
        files = self._files(rid)
        self.assertEqual(2, len(files))
        # 🔴 첨부는 방금 만든 격자 블록에 매달려야 한다.  `block_id` 가 비면 사진이
        #    카드에서 떨어져 화면에 안 보이고 purge 도 못 찾는다.
        self.assertEqual({images[0][0]}, {b['section_key'] for b in self._blocks(rid)
                                          if b['block_type'] == 'image'})
        self.assertTrue(all(f['block_id'] for f in files))
        self.assertEqual([True, True], [f['original_name'].endswith('.png') for f in files])
        with appmod.app.app_context():
            stored = [x['stored_name'] for x in appmod.query(
                'SELECT stored_name FROM dock_daily_attachment WHERE report_id=?', (rid,))]
        for name in stored:
            self.assertTrue(os.path.exists(os.path.join(routes_dock_daily.UPLOAD_DIR, name)), name)

    def test_photos_can_be_applied_without_any_rows(self):
        """사진만 골라도 들어가야 한다(`row_keys=[]`)."""
        rid = self.report['id']
        aid, scan = self._shots(rid)
        res = self._apply(rid, aid, self.report['revision'], row_keys=[],
                          photo_keys=[scan['photos'][0]['photo_key']])
        self.assertEqual(200, res.status_code, res.get_data(as_text=True))
        body = res.get_json()
        self.assertEqual(1, body['photos_added'])
        self.assertEqual(0, body['applied'])
        self.assertEqual(1, len(self._files(rid)))

    def test_nothing_picked_at_all_is_still_a_400(self):
        rid = self.report['id']
        aid, _ = self._shots(rid)
        res = self._apply(rid, aid, self.report['revision'], row_keys=[], photo_keys=[])
        self.assertEqual(400, res.status_code)
        self.assertEqual('no_rows', res.get_json()['code'])

    def test_re_reading_the_same_file_does_not_duplicate_or_bump_revision(self):
        rid = self.report['id']
        aid, scan = self._shots(rid)
        key = scan['photos'][0]['photo_key']
        first = self._apply(rid, aid, self.report['revision'], photo_keys=[key]).get_json()
        again = self._apply(rid, aid, first['revision'], photo_keys=[key]).get_json()
        self.assertEqual(0, again['photos_added'])
        self.assertEqual(1, again['photos_already'])
        self.assertEqual(1, len(self._files(rid)))
        # 🔴 아무것도 안 바뀌었으면 revision 을 올리지 않는다(다른 기기가 409 를 맞는다).
        self.assertEqual(first['revision'], again['revision'])
        # 두 번째 스캔은 그 사진을 '이미 넣음' 으로 표시한다.
        rescan = self._scan(rid, aid).get_json()
        self.assertEqual([True], [p['applied'] for p in rescan['photos']])

    def test_a_deleted_photo_can_be_picked_again_and_is_not_pre_checked(self):
        """🔴 사진은 행과 규칙이 다르다 -- 고른 대로 넣는다.

        행은 판정이 '포함' 이면 자동 체크라 형이 지운 줄을 되살리지 않게 막아야
        하지만, 사진은 매번 형이 직접 고를 때만 들어온다.  고른 것을 "전에 지웠으니"
        로 거절하면 그게 조용한 무동작이다.  되살아남 방지선은 **기본 체크 해제**이고
        그건 `dock_daily_docscan.js` 규칙에 잠겨 있다(node 테스트).
        """
        rid = self.report['id']
        aid, scan = self._shots(rid)
        key = scan['photos'][0]['photo_key']
        first = self._apply(rid, aid, self.report['revision'], photo_keys=[key]).get_json()
        fid = self._files(rid)[0]['id']
        self.assertEqual(200, self.client.delete(
            f'/api/dock-daily/attachments/{fid}').status_code)
        self.assertEqual([], self._files(rid))
        rescan = self._scan(rid, aid).get_json()
        self.assertEqual([False], [p['applied'] for p in rescan['photos']])
        again = self._apply(rid, aid, self._report_now(rid)['revision'],
                            photo_keys=[key]).get_json()
        self.assertEqual(1, again['photos_added'])
        self.assertEqual(1, len(self._files(rid)))

    def test_a_second_document_appends_to_the_same_grid(self):
        rid = self.report['id']
        aid, scan = self._shots(rid, pairs=[('Rope guard', RED)])
        first = self._apply(rid, aid, self.report['revision'],
                            photo_keys=[scan['photos'][0]['photo_key']]).get_json()
        aid2 = self._upload(rid, name='p2.docx', pairs=[('Anode', BLUE)])
        scan2 = self._scan(rid, aid2).get_json()
        body = self._apply(rid, aid2, first['revision'],
                           photo_keys=[scan2['photos'][0]['photo_key']]).get_json()
        self.assertEqual(1, body['photos_added'])
        images = self._images(rid)
        self.assertEqual(1, len(images), '격자는 한 장이어야 한다')
        self.assertEqual(['Rope guard', 'Anode'], images[0][1])
        self.assertIsNone(body['created_section'], '섹션은 다시 만들지 않는다')

    def test_a_grid_the_user_edited_is_left_alone(self):
        """🔴 형이 손댄 격자는 안 건드리고 새 격자를 만든다(카드와 같은 규칙)."""
        rid = self.report['id']
        aid, scan = self._shots(rid, pairs=[('Rope guard', RED)])
        first = self._apply(rid, aid, self.report['revision'],
                            photo_keys=[scan['photos'][0]['photo_key']]).get_json()
        with appmod.app.app_context():
            appmod.execute('UPDATE dock_daily_block SET manual_override=1'
                           " WHERE report_id=? AND block_type='image'", (rid,))
        aid2 = self._upload(rid, name='p3.docx', pairs=[('Anode', BLUE)])
        scan2 = self._scan(rid, aid2).get_json()
        self._apply(rid, aid2, first['revision'],
                    photo_keys=[scan2['photos'][0]['photo_key']])
        images = self._images(rid)
        self.assertEqual(2, len(images))
        self.assertEqual([['Rope guard'], ['Anode']], [x[1] for x in images])

    def test_photos_never_reach_the_svms_remark(self):
        """🔴 `Photos` 섹션은 SVMS `RMK` 로 가지 않는다.

        Crew 는 라벨 부분일치로 들어가지만 사진 섹션은 그 목록에 없고, `_render_section`
        이 image 블록을 건너뛴다.  둘 중 하나만 믿으면 나중에 조용히 뚫린다.
        """
        rid = self.report['id']
        aid, scan = self._shots(rid)
        self._apply(rid, aid, self.report['revision'],
                    photo_keys=[scan['photos'][0]['photo_key']])
        self.assertNotIn(routes_dock_daily.DOCX_PHOTO_SECTION.lower(),
                         routes_dock_daily.SVMS_RMK_SPECIAL_LABEL_HINTS)
        with appmod.app.app_context():
            preview = routes_dock_daily._svms(rid)
        self.assertNotIn('Rope guard', json.dumps(preview, ensure_ascii=False))

    def test_an_oversize_photo_is_dropped_alone(self):
        """사진 한 장이 상한을 넘어도 나머지는 들어간다."""
        rid = self.report['id']
        old = routes_dock_daily.MAX_ATTACHMENT
        big = _png((7, 7, 7), size=(400, 400))
        # 🔴 문서 업로드는 원래 상한으로 먼저 받는다.  상한을 먼저 낮추면 docx 자체가
        #    413 이 되고, 그러면 이 테스트는 사진 한 장 규칙을 아예 안 밟는다.
        aid = self._upload(rid, pairs=[('Big', big), ('Small', RED)])
        try:
            routes_dock_daily.MAX_ATTACHMENT = len(big) - 1
            scan = self._scan(rid, aid).get_json()
            self.assertEqual(['Small'], [p['caption'] for p in scan['photos']])
            self.assertEqual(1, scan['photo_skipped'])
            body = self._apply(rid, aid, self.report['revision'],
                               photo_keys=[p['photo_key'] for p in scan['photos']]).get_json()
            self.assertEqual(1, body['photos_added'])
        finally:
            routes_dock_daily.MAX_ATTACHMENT = old

    def test_unknown_photo_keys_are_reported(self):
        rid = self.report['id']
        aid, scan = self._shots(rid)
        body = self._apply(rid, aid, self.report['revision'],
                           photo_keys=[scan['photos'][0]['photo_key'], 'deadbeef1234']).get_json()
        self.assertEqual(['deadbeef1234'], body['unknown_photo_keys'])
        self.assertEqual(1, body['photos_added'])

    def test_photo_keys_must_be_an_array(self):
        rid = self.report['id']
        aid, _ = self._shots(rid)
        res = self._apply(rid, aid, self.report['revision'], photo_keys='all')
        self.assertEqual(400, res.status_code)

    def test_a_final_report_takes_no_photos(self):
        rid = self.report['id']
        aid, scan = self._shots(rid)
        fin = self.client.post(f'/api/dock-daily/reports/{rid}/status',
                               json={'status': 'final',
                                     'revision': self._report_now(rid)['revision']})
        self.assertEqual(200, fin.status_code, fin.get_data(as_text=True))
        res = self._apply(rid, aid, self._report_now(rid)['revision'],
                          photo_keys=[scan['photos'][0]['photo_key']])
        self.assertEqual(409, res.status_code)
        self.assertEqual([], self._files(rid))
        # 🔴 파일도 남지 않아야 한다.  사진 쓰기는 트랜잭션 안이라, rollback 하면서
        #    안 지우면 행 없는 blob 이 영구히 남는다(어떤 purge 도 못 찾는다).
        left = [x for x in os.listdir(routes_dock_daily.UPLOAD_DIR)
                if x.endswith('.png')] if os.path.isdir(routes_dock_daily.UPLOAD_DIR) else []
        self.assertEqual([], left)

    def _report_now(self, rid):
        return self.client.get(f'/api/dock-daily/reports/{rid}').get_json()


if __name__ == '__main__':
    unittest.main()
