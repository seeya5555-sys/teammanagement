import hashlib
import io
import os
import sqlite3
import tempfile
import unittest

from docx import Document

import app as appmod
import ai_gemini as routes


class SvmsSireAttachmentTests(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.old_db = appmod.DATABASE
        self.old_cfg = appmod.app.config["DATABASE"]
        self.old_upload = routes.UPLOAD_DIR
        db = os.path.join(self.tmp.name, "svms-sire.db")
        appmod.DATABASE = db
        appmod.app.config["DATABASE"] = db
        routes.UPLOAD_DIR = os.path.join(self.tmp.name, "uploads")
        os.makedirs(routes.UPLOAD_DIR)
        with appmod.app.app_context():
            appmod.init_db(False)
            appmod._auto_migrate()
            appmod.execute("INSERT OR REPLACE INTO api_settings(k,v) VALUES('api_key',?)", ("test-key",))
            vessel = appmod.execute(
                "INSERT INTO vessels(name,short_name,active) VALUES(?,?,1)",
                ("KUWAIT PROSPERITY", "KUWAIT"),
            )
            self.vetting = appmod.execute(
                "INSERT INTO vettings(vessel_id,report_number) VALUES(?,?)",
                (vessel, "LZXN-1955-2760-7793"),
            )
        self.client = appmod.app.test_client()

    def tearDown(self):
        routes.UPLOAD_DIR = self.old_upload
        appmod.DATABASE = self.old_db
        appmod.app.config["DATABASE"] = self.old_cfg
        self.tmp.cleanup()

    def _post(self, payload=b"%PDF-1.4\nreport\n%%EOF", **overrides):
        fields = {
            "vessel_name": "kuwait-prosperity",
            "report_number": "LZXN195527607793",
            "SIRE_CD": "KWPSSR2606020008",
            "FILE_TP": "SMSC",
            "external_file_id": "KWPSSR2606020008|SMSC|SKSMSMSC2606260003",
            "sha256": hashlib.sha256(payload).hexdigest(),
        }
        fields.update(overrides)
        fields["file"] = (io.BytesIO(payload), "close.pdf")
        return self.client.post(
            "/api/ext/vettings/svms-attachment",
            data=fields,
            content_type="multipart/form-data",
            headers={"X-API-Key": "test-key"},
        )

    def test_upload_deduplicates_and_preserves_findings(self):
        first = self._post()
        self.assertEqual(201, first.status_code, first.get_data(as_text=True))
        second = self._post()
        self.assertEqual(200, second.status_code, second.get_data(as_text=True))
        self.assertTrue(second.get_json()["deduplicated"])
        with appmod.app.app_context():
            rows = appmod.query("SELECT * FROM vt_attachments WHERE vetting_id=?", (self.vetting,))
            self.assertEqual(1, len(rows))
            self.assertEqual("svms", rows[0]["source"])
            self.assertEqual("close", rows[0]["source_type"])
            self.assertEqual(0, appmod.query("SELECT COUNT(*) n FROM vt_findings", one=True)["n"])
        self.assertEqual(1, len(os.listdir(routes.UPLOAD_DIR)))

    def test_rejects_bad_magic_and_ambiguous_match(self):
        bad = self._post(b"not a pdf")
        self.assertEqual(415, bad.status_code)
        with appmod.app.app_context():
            vessel = appmod.execute(
                "INSERT INTO vessels(name,short_name,active) VALUES(?,?,1)",
                ("KUWAIT-PROSPERITY", "KUWAIT2"),
            )
            appmod.execute(
                "INSERT INTO vettings(vessel_id,report_number) VALUES(?,?)",
                (vessel, "LZXN195527607793"),
            )
        ambiguous = self._post()
        self.assertEqual(409, ambiguous.status_code)
        self.assertIn("ambiguous", ambiguous.get_json()["error"])
        self.assertEqual([], os.listdir(routes.UPLOAD_DIR))

    def test_svms_attachment_cannot_be_deleted(self):
        uploaded = self._post().get_json()
        with self.client.session_transaction() as session:
            session.update(user_id=1, username="admin", role="admin", supervisor_id=None)
        response = self.client.delete(f"/api/vt-attachments/{uploaded['id']}")
        self.assertEqual(403, response.status_code)

    def test_new_hash_supersedes_old_revision_and_list_returns_provenance(self):
        old = self._post().get_json()["id"]
        new = self._post(b"%PDF-1.4\nrevised\n%%EOF").get_json()["id"]
        self.assertNotEqual(old, new)
        with appmod.app.app_context():
            rows = appmod.query(
                "SELECT id,inactive_at FROM vt_attachments WHERE external_file_id=? ORDER BY id",
                ("KWPSSR2606020008|SMSC|SKSMSMSC2606260003",),
            )
            self.assertEqual(2, len(rows))
            self.assertIsNotNone(rows[0]["inactive_at"])
            self.assertIsNone(rows[1]["inactive_at"])
        with self.client.session_transaction() as session:
            session.update(user_id=1, username="admin", role="admin", supervisor_id=None)
        listed = self.client.get(f"/api/vettings/{self.vetting}/attachments")
        self.assertEqual(200, listed.status_code)
        payload = listed.get_json()
        self.assertEqual([new], [row["id"] for row in payload])
        self.assertEqual("svms", payload[0]["source"])
        self.assertEqual("close", payload[0]["source_type"])
        self.assertTrue(payload[0]["synced_at"])

    def test_smsr_is_stored_as_initial(self):
        uploaded = self._post(
            FILE_TP="SMSR",
            external_file_id="KWPSSR2606020008|SMSR|SKSMSMSR2606020004",
        )
        self.assertEqual(201, uploaded.status_code, uploaded.get_data(as_text=True))
        with appmod.app.app_context():
            row = appmod.query("SELECT source_type FROM vt_attachments WHERE id=?",
                               (uploaded.get_json()["id"],), one=True)
            self.assertEqual("initial", row["source_type"])

    def test_api_key_and_exact_match_are_required(self):
        anon = self.client.post("/api/ext/vettings/svms-attachment")
        self.assertEqual(401, anon.status_code)
        missing = self._post(vessel_name="OTHER VESSEL")
        self.assertEqual(409, missing.status_code)

    def test_status_flags_update_exact_report_without_touching_findings(self):
        response = self.client.post(
            "/api/ext/vettings/svms-status",
            json={"vessel_name":"kuwait-prosperity", "report_number":"LZXN195527607793",
                  "full_report_yn":"Y", "close_report_yn":"N"},
            headers={"X-API-Key":"test-key"},
        )
        self.assertEqual(200, response.status_code, response.get_data(as_text=True))
        with appmod.app.app_context():
            row = appmod.query("SELECT * FROM vettings WHERE id=?", (self.vetting,), one=True)
            self.assertEqual("Y", row["svms_full_report_yn"])
            self.assertEqual("N", row["svms_close_report_yn"])
            self.assertEqual("Y", row["svms_report_uploaded_yn"])
            self.assertTrue(row["svms_status_synced_at"])
            self.assertEqual(0, appmod.query("SELECT COUNT(*) n FROM vt_findings", one=True)["n"])
        listed = self.client.get("/api/ext/vettings", headers={"X-API-Key":"test-key"}).get_json()
        payload = next(v for v in listed if v["id"] == self.vetting)
        self.assertEqual("Y", payload["svms_full_report_yn"])
        self.assertEqual("N", payload["svms_close_report_yn"])
        self.assertEqual("Y", payload["svms_report_uploaded_yn"])
        bad = self.client.post(
            "/api/ext/vettings/svms-status",
            json={"vessel_name":"KUWAIT PROSPERITY", "report_number":"LZXN195527607793",
                  "full_report_yn":"?", "close_report_yn":"N"},
            headers={"X-API-Key":"test-key"},
        )
        self.assertEqual(400, bad.status_code)
        missing = self.client.post(
            "/api/ext/vettings/svms-status",
            json={"vessel_name":"OTHER", "report_number":"NONE",
                  "full_report_yn":"N", "close_report_yn":"N"},
            headers={"X-API-Key":"test-key"},
        )
        self.assertEqual(409, missing.status_code)

    def test_status_marks_report_not_uploaded_and_clears_stale_flags(self):
        with appmod.app.app_context():
            appmod.execute(
                "UPDATE vettings SET svms_full_report_yn='Y', svms_close_report_yn='N' WHERE id=?",
                (self.vetting,),
            )
        response = self.client.post(
            "/api/ext/vettings/svms-status",
            json={"vessel_name":"KUWAIT PROSPERITY", "report_number":"LZXN195527607793",
                  "report_uploaded_yn":"N"},
            headers={"X-API-Key":"test-key"},
        )
        self.assertEqual(200, response.status_code, response.get_data(as_text=True))
        with appmod.app.app_context():
            row = appmod.query("SELECT * FROM vettings WHERE id=?", (self.vetting,), one=True)
        self.assertEqual("N", row["svms_report_uploaded_yn"])
        self.assertIsNone(row["svms_full_report_yn"])
        self.assertIsNone(row["svms_close_report_yn"])
        self.assertTrue(row["svms_status_synced_at"])

    def test_docx_opens_as_safe_inline_html_and_can_download_original(self):
        stored = "preview.docx"
        path = os.path.join(routes.UPLOAD_DIR, stored)
        doc = Document()
        doc.add_heading("Close Report", level=1)
        doc.add_paragraph("<script>alert(1)</script>")
        table = doc.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Observation"
        table.cell(0, 1).text = "Closed"
        doc.save(path)
        with appmod.app.app_context():
            aid = appmod.execute(
                "INSERT INTO vt_attachments(vetting_id,filename,stored_name,file_size,mime_type) VALUES(?,?,?,?,?)",
                (self.vetting, "close report.docx", stored, os.path.getsize(path),
                 "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            )
        with self.client.session_transaction() as session:
            session.update(user_id=1, username="admin", role="admin", supervisor_id=None)
        preview = self.client.get(f"/api/vt-attachments/{aid}/docx-preview")
        self.assertEqual(200, preview.status_code, preview.get_data(as_text=True))
        html = preview.get_data(as_text=True)
        self.assertIn("Close Report", html)
        self.assertIn("Observation", html)
        self.assertIn("&lt;script&gt;alert(1)&lt;/script&gt;", html)
        self.assertNotIn("<script>alert(1)</script>", html)
        download = self.client.get(f"/api/vt-attachments/{aid}/docx-preview?download=1")
        self.assertEqual(200, download.status_code)
        self.assertIn("attachment", download.headers.get("Content-Disposition", ""))
        download.close()

    def test_legacy_database_migrates_before_unique_index_creation(self):
        legacy = os.path.join(self.tmp.name, "legacy.db")
        conn = sqlite3.connect(legacy)
        conn.execute("""CREATE TABLE vt_attachments(
            id INTEGER PRIMARY KEY, vetting_id INTEGER, filename TEXT,
            stored_name TEXT, file_size INTEGER, mime_type TEXT,
            uploaded_by TEXT, uploaded_at TEXT)""")
        conn.commit(); conn.close()
        with appmod.app.app_context():
            appmod.DATABASE = legacy
            appmod.app.config["DATABASE"] = legacy
            appmod.init_db(False)
            appmod._auto_migrate()
            cols = {row[1] for row in appmod.get_db().execute("PRAGMA table_info(vt_attachments)")}
            indexes = {row[1] for row in appmod.get_db().execute("PRAGMA index_list(vt_attachments)")}
        self.assertIn("external_file_id", cols)
        self.assertIn("uq_vt_attachments_svms_identity_sha", indexes)


if __name__ == "__main__":
    unittest.main()
