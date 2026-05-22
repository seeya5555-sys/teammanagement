"""
Boarding Report — Word(.docx) 생성

방선보고서 + Defect List 통합 양식:
  · 표지: "Vessel boarding report" 헤더 + 결재란 + 정보 표 + Master/CE
  · 본문: 섹션별 블록 (paragraph / bullet / table / image / info_table / defect_table)
  · Sinokor 푸터: "CODE<107-301>/2015.04.17 ... Sinokor Ship Management Co., Ltd"

Dry Dock 모듈(dock_report_docx.py)의 헬퍼 일부를 임포트해 재사용.
"""
import io
import json
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Dry Dock 모듈에서 검증된 헬퍼 재사용
from dock_report_docx import (
    _set_cell_shading, _set_cell_borders, _set_font, _add_paragraph,
    _set_row_height, _set_table_fixed_layout, _add_horizontal_line,
    _crop_to_aspect, _GLOBAL_TEMP_FILES,
)


# ─────────────────────────────────────────────────────────────
#  방선보고서 표지 — "Vessel boarding report" 헤더 + 결재란 + 정보 표
# ─────────────────────────────────────────────────────────────
def _build_brep_cover(doc, report):
    """첨부 양식의 첫 페이지 (헤더 + 결재란 + 정보 표)"""

    # ① 상단 큰 타이틀
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Vessel Boarding Report')
    _set_font(r, name='Malgun Gothic', size=22, bold=True, color='1F4E79')

    # 밑줄
    line_p = doc.add_paragraph()
    pPr = line_p._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '18')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F4E79')
    pBdr.append(bottom)
    pPr.append(pBdr)

    _add_paragraph(doc, '', before=8)

    # ② 결재란 (우측 정렬)
    approvals = [
        ('기 안 자', report.get('approval_drafter') or ''),
        ('팀 장',    report.get('approval_team_lead') or ''),
        ('중 역',    report.get('approval_director') or ''),
        ('대표이사', report.get('approval_ceo') or ''),
    ]
    app_tbl = doc.add_table(rows=2, cols=len(approvals) + 1)
    app_tbl.alignment = WD_TABLE_ALIGNMENT.RIGHT
    app_tbl.autofit = False

    # 표 자체를 우측으로 밀어붙임
    tblPr = app_tbl._element.find(qn('w:tblPr'))
    if tblPr is not None:
        existing = tblPr.find(qn('w:tblInd'))
        if existing is not None: tblPr.remove(existing)
        tblInd = OxmlElement('w:tblInd')
        tblInd.set(qn('w:w'), '4250')  # ~7.5cm 들여쓰기
        tblInd.set(qn('w:type'), 'dxa')
        tblPr.append(tblInd)

    # "결재" 라벨 셀 (세로 병합)
    title_cell = app_tbl.rows[0].cells[0]
    title_cell.merge(app_tbl.rows[1].cells[0])
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _set_cell_shading(title_cell, 'F2F2F2')
    tp = title_cell.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run('결재')
    _set_font(tr, size=11, bold=True)
    _set_cell_borders(title_cell)

    # 결재자 헤더 + 서명란
    for col, (label, name) in enumerate(approvals, start=1):
        h = app_tbl.rows[0].cells[col]
        _set_cell_shading(h, 'F2F2F2')
        h.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hp = h.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(label)
        _set_font(hr, size=9, bold=True)
        _set_cell_borders(h)

        sig = app_tbl.rows[1].cells[col]
        sig.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        sp = sig.paragraphs[0]
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if name:
            sr = sp.add_run(name)
            _set_font(sr, size=10)
        _set_cell_borders(sig)
    _set_row_height(app_tbl.rows[1], 1.3)

    _add_paragraph(doc, '', before=12)

    # ③ 정보 표 (양식의 4행 헤더 표 — Vessel / Port / Inspector / Date 등)
    info_data = [
        ('Vessel',    report.get('vessel_name') or '',
         'Inspector', report.get('supervisor_name') or ''),
        ('Port',      report.get('port') or '',
         'Date / Time', _fmt_period(report.get('boarding_start'),
                                     report.get('boarding_end'))),
        ('Master (boarding date)',
            _fmt_person_date(report.get('master_name'),
                             report.get('master_board_date')),
         'C/E (boarding date)',
            _fmt_person_date(report.get('chief_eng_name'),
                             report.get('chief_eng_board_date'))),
        ('Ship-Visit Checklist Score',
            report.get('sv_checklist_score') or '-',
         'Reported on',
            datetime.now().strftime('%Y-%m-%d')),
    ]

    info_tbl = doc.add_table(rows=len(info_data), cols=4)
    info_tbl.autofit = False
    col_cm = [3.8, 5.2, 3.8, 5.2]   # 합계 18cm 정도

    for ri, row in enumerate(info_data):
        for ci, val in enumerate(row):
            cell = info_tbl.rows[ri].cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            _set_cell_borders(cell)
            p = cell.paragraphs[0]
            # 라벨 셀(0, 2): 음영 + 굵게 + 가운데
            if ci % 2 == 0:
                _set_cell_shading(cell, 'E7EBF0')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(str(val))
                _set_font(r, size=10, bold=True)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                r = p.add_run(str(val))
                _set_font(r, size=10)

    _set_table_fixed_layout(info_tbl, 18.0, col_cm)

    _add_paragraph(doc, '', before=12)


def _fmt_period(start, end):
    if not start and not end:
        return ''
    s = (start or '').replace('-', '.')
    e = (end or '').replace('-', '.')
    if s and e:
        return f'{s} ~ {e}'
    return s or e


def _fmt_person_date(name, date):
    name = (name or '').strip()
    date = (date or '').replace('-', '.')
    if name and date:
        return f'{name}  ({date})'
    return name or date or ''


# ─────────────────────────────────────────────────────────────
#  본문 — 섹션 + 블록
# ─────────────────────────────────────────────────────────────
def _render_brep_sections(doc, sections_tree, depth=0, prefix=''):
    for i, sec in enumerate(sections_tree):
        num = f'{prefix}-{i + 1}' if prefix else f'{i + 1}'

        # 1단계 섹션 사이엔 페이지 분리
        if depth == 0 and i > 0:
            doc.add_page_break()

        if depth == 0:
            _add_paragraph(doc, f'{num}. {sec["title"]}',
                           size=14, bold=True, color='1F4E79',
                           before=14, after=6)
            _add_horizontal_line(doc, color='1F4E79', sz=8)
        elif depth == 1:
            _add_paragraph(doc, f'{num}. {sec["title"]}',
                           size=13, bold=True, color='1F4E79',
                           before=10, after=4)
            _add_horizontal_line(doc, color='5B9BD5', sz=6)
        else:
            _add_paragraph(doc, f'{num}. {sec["title"]}',
                           size=11.5, bold=True, color='2E5990',
                           before=8, after=4)

        blocks = sorted(sec.get('blocks', []),
                        key=lambda b: (b.get('display_order', 0), b.get('id', 0)))
        for b in blocks:
            _render_brep_block(doc, b, depth)

        if sec.get('children'):
            _render_brep_sections(doc, sec['children'], depth + 1, num)


def _render_brep_block(doc, block, depth):
    bt = block.get('block_type')
    content = block.get('content') or {}
    if isinstance(content, str):
        try: content = json.loads(content)
        except Exception: content = {}

    base_indent = 0.3 + 0.3 * depth

    if bt == 'paragraph':
        _render_paragraph(doc, content, base_indent)
    elif bt == 'bullet_list':
        _render_bullet(doc, content, base_indent)
    elif bt == 'table':
        _render_table(doc, content, base_indent)
    elif bt == 'image':
        _render_image(doc, content, base_indent)
    elif bt == 'info_table':
        _render_info_table(doc, content, base_indent)
    elif bt == 'defect_table':
        _render_defect_table(doc, content, base_indent)


def _render_paragraph(doc, content, base_indent):
    text = (content.get('text') or '').strip()
    if not text:
        return
    for line in text.split('\n'):
        _add_paragraph(doc, line, size=10.5,
                       indent_left=base_indent, after=4)


# 마커 형식: 깊이별
CIRCLED = ['①','②','③','④','⑤','⑥','⑦','⑧','⑨','⑩',
           '⑪','⑫','⑬','⑭','⑮','⑯','⑰','⑱','⑲','⑳']

def _alpha_char(n): return chr(96 + ((n - 1) % 26) + 1)
def _circled(n):    return CIRCLED[(n - 1) % len(CIRCLED)]

def _number_by_depth(d, n):
    if d == 0: return f'{n}.'
    if d == 1: return f'{n})'
    if d == 2: return _circled(n)
    return f'{_alpha_char(n)})'

def _alpha_by_depth(d, n):
    if d == 0: return f'{_alpha_char(n)}.'
    if d == 1: return f'{_alpha_char(n)})'
    if d == 2: return _circled(n)
    return f'{n})'


def _render_bullet(doc, content, base_indent):
    items = content.get('items') or []
    marker = content.get('marker') or 'bullet'
    counters = [0, 0, 0, 0]
    for it in items:
        if isinstance(it, str):
            text, indent = it, 0
        else:
            text = it.get('text', '')
            indent = max(0, min(3, it.get('indent', 0)))
        if not text.strip():
            continue
        counters[indent] += 1
        for k in range(indent + 1, 4):
            counters[k] = 0
        n = counters[indent]
        if marker == 'dash':      mk = '–'
        elif marker == 'number':  mk = _number_by_depth(indent, n)
        elif marker == 'alpha':   mk = _alpha_by_depth(indent, n)
        else:                      mk = '•'
        _add_paragraph(doc, f'{mk}  {text}', size=10.5,
                       indent_left=base_indent + 0.6 * indent, after=2)


def _render_table(doc, content, base_indent):
    headers = content.get('headers') or []
    rows = content.get('rows') or []
    if not headers and not rows:
        return
    n_cols = max(len(headers), max((len(r) for r in rows), default=0))
    if n_cols == 0:
        return
    headers = (list(headers) + [''] * n_cols)[:n_cols]
    rows = [(list(r) + [''] * n_cols)[:n_cols] for r in rows]

    col_widths_px = content.get('col_widths') or []
    total_cm = 16.0
    if (not col_widths_px or len(col_widths_px) != n_cols
            or sum(w for w in col_widths_px if w and w > 0) <= 0):
        col_cm = [total_cm / n_cols] * n_cols
    else:
        valid = [w for w in col_widths_px if w and w > 0]
        if len(valid) < n_cols:
            avg = sum(valid) / len(valid)
            col_widths_px = [w if (w and w > 0) else avg for w in col_widths_px]
        total = sum(col_widths_px)
        col_cm = [total_cm * (w / total) for w in col_widths_px]

    tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for ci, h in enumerate(headers):
        cell = tbl.rows[0].cells[ci]
        _set_cell_shading(cell, 'D9E2EC')
        _set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        _set_font(r, size=10, bold=True)

    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            _set_cell_borders(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            lines = str(val or '').split('\n')
            for li, ln in enumerate(lines):
                if li > 0:
                    p = cell.add_paragraph()
                r = p.add_run(ln)
                _set_font(r, size=10)

    _set_table_fixed_layout(tbl, total_cm, col_cm)
    _add_paragraph(doc, '', before=2, after=4)


def _render_image(doc, content, base_indent):
    images = content.get('images') or []
    columns = max(1, min(4, int(content.get('columns', 2) or 2)))
    if not images:
        return
    n_rows = (len(images) + columns - 1) // columns
    total_cm = 16.0
    cell_cm = (total_cm - 0.3 * (columns - 1)) / columns
    img_w = cell_cm - 0.4
    img_h = img_w * 3 / 4

    tbl = doc.add_table(rows=n_rows * 2, cols=columns)
    tbl.autofit = False
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for idx, img in enumerate(images):
        ri = (idx // columns) * 2
        ci = idx % columns
        img_cell = tbl.rows[ri].cells[ci]
        cap_cell = tbl.rows[ri + 1].cells[ci]

        img_path = _resolve_brep_image_path(img.get('url') or '',
                                             img.get('filename') or '')
        if img_path and os.path.exists(img_path):
            try:
                processed = _crop_to_aspect(img_path, 4/3)
                if processed != img_path:
                    _GLOBAL_TEMP_FILES.append(processed)
                p = img_cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(processed, width=Cm(img_w), height=Cm(img_h))
            except Exception as e:
                p = img_cell.paragraphs[0]
                rr = p.add_run(f'[이미지 로드 실패: {e}]')
                _set_font(rr, size=9, color='B91C1C')
        else:
            p = img_cell.paragraphs[0]
            rr = p.add_run('[이미지 없음]')
            _set_font(rr, size=9, color='9CA3AF')

        cap_p = cap_cell.paragraphs[0]
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = img.get('caption') or ''
        if caption:
            cr = cap_p.add_run(caption)
            _set_font(cr, size=9, color='4B5563')
            cr.italic = True

    col_cm_list = [cell_cm] * columns
    _set_table_fixed_layout(tbl, total_cm, col_cm_list)
    _add_paragraph(doc, '', before=2, after=6)


def _resolve_brep_image_path(url, filename):
    """boarding/ 폴더에 저장된 이미지 경로 해결"""
    if url and url.startswith('/static/'):
        rel = url[len('/static/'):]
        from app import app
        return os.path.join(app.static_folder, rel)
    if filename:
        from app import app
        return os.path.join(app.static_folder, 'uploads', 'boarding', filename)
    return None


# ─────────────────────────────────────────────────────────────
#  신규 블록: info_table (Label-Value 표)
# ─────────────────────────────────────────────────────────────
def _render_info_table(doc, content, base_indent):
    rows = content.get('rows') or []
    rows = [r for r in rows if (r.get('label') or '').strip() or (r.get('value') or '').strip()]
    if not rows:
        return

    total_cm = 16.0
    col_cm = [4.5, total_cm - 4.5]

    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.autofit = False
    for ri, r in enumerate(rows):
        # Label 셀
        lcell = tbl.rows[ri].cells[0]
        _set_cell_shading(lcell, 'F3F4F6')
        _set_cell_borders(lcell)
        lcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        lp = lcell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lr = lp.add_run(r.get('label') or '')
        _set_font(lr, size=10.5, bold=True)

        # Value 셀
        vcell = tbl.rows[ri].cells[1]
        _set_cell_borders(vcell)
        vcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        vp = vcell.paragraphs[0]
        vp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # 줄바꿈 보존
        lines = (r.get('value') or '').split('\n')
        for li, ln in enumerate(lines):
            if li > 0:
                vp = vcell.add_paragraph()
            vr = vp.add_run(ln)
            _set_font(vr, size=10.5)

    _set_table_fixed_layout(tbl, total_cm, col_cm)
    _add_paragraph(doc, '', before=2, after=4)


# ─────────────────────────────────────────────────────────────
#  신규 블록: defect_table (Defect List)
# ─────────────────────────────────────────────────────────────
RISK_COLORS = {
    'L': {'bg': 'D1FAE5', 'fg': '065F46'},   # 초록
    'M': {'bg': 'FEF3C7', 'fg': '92400E'},   # 노랑
    'H': {'bg': 'FEE2E2', 'fg': '991B1B'},   # 빨강
}

def _render_defect_table(doc, content, base_indent):
    items = content.get('items') or []
    if not items:
        return

    # 헤더 + 데이터 행
    total_cm = 16.0
    # 컬럼: No(0.8) / Photo(4) / Description(5.6) / Rectification(5.6)
    col_cm = [0.8, 4.0, 5.6, 5.6]
    n_cols = 4

    tbl = doc.add_table(rows=1 + len(items), cols=n_cols)
    tbl.autofit = False

    # 헤더
    headers = ['', 'Item (Photo)', 'Description (Findings)', 'Rectification']
    for ci, h in enumerate(headers):
        cell = tbl.rows[0].cells[ci]
        _set_cell_shading(cell, '1F4E79')
        _set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        _set_font(r, size=10.5, bold=True, color='FFFFFF')

    # 데이터 행
    for idx, it in enumerate(items, start=1):
        risk = (it.get('risk') or 'L').upper()
        if risk not in ('L', 'M', 'H'):
            risk = 'L'
        row_bg = RISK_COLORS[risk]['bg']

        cells = tbl.rows[idx].cells

        # No 셀
        no_cell = cells[0]
        _set_cell_shading(no_cell, row_bg)
        _set_cell_borders(no_cell)
        no_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        np_ = no_cell.paragraphs[0]
        np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nr = np_.add_run(str(idx))
        _set_font(nr, size=11, bold=True)

        # Photo 셀
        ph_cell = cells[1]
        _set_cell_shading(ph_cell, row_bg)
        _set_cell_borders(ph_cell)
        ph_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _render_defect_photos(ph_cell, it.get('images') or [], col_cm[1])

        # Description (Item + (Risk) + 줄바꿈 desc)
        desc_cell = cells[2]
        _set_cell_shading(desc_cell, row_bg)
        _set_cell_borders(desc_cell)
        desc_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        # 첫 줄: Item 굵게 + (Risk)
        item_text = (it.get('item') or '').strip()
        first_p = desc_cell.paragraphs[0]
        first_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if item_text:
            run = first_p.add_run(item_text)
            _set_font(run, size=10.5, bold=True)
            sp = first_p.add_run(f'  ({risk})')
            _set_font(sp, size=10.5, bold=True, color=RISK_COLORS[risk]['fg'])
        else:
            sp = first_p.add_run(f'({risk})')
            _set_font(sp, size=10.5, bold=True, color=RISK_COLORS[risk]['fg'])
        # 나머지 줄: desc
        desc_text = (it.get('desc') or '').strip()
        if desc_text:
            for line in desc_text.split('\n'):
                lp = desc_cell.add_paragraph()
                lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
                lr = lp.add_run(line)
                _set_font(lr, size=10)

        # Rectification
        fix_cell = cells[3]
        _set_cell_shading(fix_cell, row_bg)
        _set_cell_borders(fix_cell)
        fix_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        fix_text = (it.get('fix') or '').strip()
        fp = fix_cell.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lines = fix_text.split('\n') if fix_text else ['']
        for li, line in enumerate(lines):
            if li > 0:
                fp = fix_cell.add_paragraph()
                fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fr = fp.add_run(line)
            _set_font(fr, size=10)

    _set_table_fixed_layout(tbl, total_cm, col_cm)
    _add_paragraph(doc, '', before=2, after=2)

    # Risk Legend
    legend = doc.add_paragraph()
    legend.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr = legend.add_run('Level of Risk:    ')
    _set_font(lr, size=10, bold=True)

    for v, label, color in [('L', 'L : Low', '065F46'),
                             ('M', 'M : Medium', '92400E'),
                             ('H', 'H : High', '991B1B')]:
        rr = legend.add_run(label)
        _set_font(rr, size=10, bold=True, color=color)
        sep = legend.add_run('     ')
        _set_font(sep, size=10)

    _add_paragraph(doc, '', before=2, after=4)


def _render_defect_photos(cell, images, cell_cm):
    """defect 항목의 사진 — 셀 안에 메인 1장 + 작은 thumb 그리드"""
    if not images:
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('—')
        _set_font(r, size=10, color='9CA3AF')
        return

    main_img = images[0]
    extras = images[1:]
    img_w = cell_cm - 0.4

    # 메인 이미지
    main_path = _resolve_brep_image_path(main_img.get('url') or '',
                                          main_img.get('filename') or '')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if main_path and os.path.exists(main_path):
        try:
            processed = _crop_to_aspect(main_path, 4/3)
            if processed != main_path:
                _GLOBAL_TEMP_FILES.append(processed)
            run = p.add_run()
            run.add_picture(processed,
                            width=Cm(img_w),
                            height=Cm(img_w * 3 / 4))
        except Exception:
            r = p.add_run('[이미지 오류]')
            _set_font(r, size=8, color='B91C1C')
    else:
        r = p.add_run('[이미지 없음]')
        _set_font(r, size=8, color='9CA3AF')

    # 추가 이미지들
    if extras:
        extra_p = cell.add_paragraph()
        extra_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        thumb_w = (img_w - 0.2 * min(3, len(extras))) / min(4, len(extras))
        thumb_h = thumb_w * 3 / 4
        for ex in extras[:4]:
            ex_path = _resolve_brep_image_path(ex.get('url') or '',
                                                ex.get('filename') or '')
            if ex_path and os.path.exists(ex_path):
                try:
                    processed = _crop_to_aspect(ex_path, 4/3)
                    if processed != ex_path:
                        _GLOBAL_TEMP_FILES.append(processed)
                    run = extra_p.add_run()
                    run.add_picture(processed,
                                    width=Cm(thumb_w), height=Cm(thumb_h))
                    extra_p.add_run(' ')
                except Exception:
                    pass


# ─────────────────────────────────────────────────────────────
#  공개 함수 — 보고서 데이터 받아 docx 바이트 반환
# ─────────────────────────────────────────────────────────────
def build_docx(report: dict) -> bytes:
    doc = Document()

    # 페이지 설정
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width  = Mm(210)
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
    rFonts.set(qn('w:ascii'),    'Malgun Gothic')
    rFonts.set(qn('w:hAnsi'),    'Malgun Gothic')

    # Sinokor 푸터 (모든 페이지에 표시)
    _add_sinokor_footer(doc)

    # 섹션 트리 빌드
    sections_flat = report.get('sections') or []
    tree = _build_tree(sections_flat)

    # ① 표지 (헤더 + 결재란 + 정보 표)
    _build_brep_cover(doc, report)

    # ② 본문
    _render_brep_sections(doc, tree)

    # 직렬화
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    result = bio.read()

    # 임시 파일 정리
    global _GLOBAL_TEMP_FILES
    for fp in _GLOBAL_TEMP_FILES:
        try: os.remove(fp)
        except Exception: pass
    _GLOBAL_TEMP_FILES.clear()

    return result


def _build_tree(sections_flat):
    by_id = {s['id']: dict(s, children=[]) for s in sections_flat}
    roots = []
    for s in by_id.values():
        if s.get('parent_id') and s['parent_id'] in by_id:
            by_id[s['parent_id']]['children'].append(s)
        else:
            roots.append(s)
    def sort_rec(lst):
        lst.sort(key=lambda x: (x.get('display_order', 0), x.get('id', 0)))
        for x in lst:
            sort_rec(x['children'])
    sort_rec(roots)
    return roots


def _add_sinokor_footer(doc):
    """첨부 양식의 푸터:
       'CODE<107-301>/2015.04.17     Sinokor Ship Management Co., Ltd'
       좌측: CODE / 우측: Sinokor — 탭으로 양쪽 정렬
    """
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0]

        # 탭 스톱: 우측 정렬 위치 (본문 가용폭 16cm 기준)
        pPr = p._element.get_or_add_pPr()
        existing_tabs = pPr.find(qn('w:tabs'))
        if existing_tabs is not None:
            pPr.remove(existing_tabs)
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'), 'right')
        tab.set(qn('w:pos'), '9355')   # ~16.5cm in twips
        tabs.append(tab)
        pPr.append(tabs)

        r1 = p.add_run('CODE<107-301>/2015.04.17')
        _set_font(r1, size=8, color='6B7280')

        p.add_run('\t')

        r2 = p.add_run('Sinokor Ship Management Co., Ltd')
        _set_font(r2, size=8, color='6B7280', bold=True)
